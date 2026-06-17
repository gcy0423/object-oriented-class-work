from __future__ import annotations

import os
import textwrap
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "EduMindAgent_结课设计文档.docx"
FONT_CJK = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(23, 32, 51)
MUTED = RGBColor(102, 112, 133)
LIGHT_FILL = "F2F4F7"
SOFT_BLUE = "E8EEF5"
WHITE = "FFFFFF"
LINE = "D8DEE9"


def font_path() -> str:
    for candidate in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if os.path.exists(candidate):
            return candidate
    return ""


def pil_font(size: int, bold: bool = False):
    path = font_path()
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap_draw_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        trial = current + char
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill="#172033", line_gap=5):
    x1, y1, x2, y2 = box
    lines = wrap_draw_text(draw, text, font, int(x2 - x1 - 24))
    total_h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += (bbox[3] - bbox[1]) + line_gap


def box(draw, xy, text, fill="#FFFFFF", outline="#B8C2D6", width=2, font_size=28, radius=16):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)
    draw_centered_text(draw, xy, text, pil_font(font_size), fill="#172033")


def arrow(draw, start, end, fill="#52637A", width=3):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    dx = 1 if ex >= sx else -1
    draw.polygon([(ex, ey), (ex - 16 * dx, ey - 8), (ex - 16 * dx, ey + 8)], fill=fill)


def save_feature_tree(path: Path):
    img = Image.new("RGB", (1800, 1120), "#F7F8FB")
    draw = ImageDraw.Draw(img)
    title = pil_font(42)
    draw.text((60, 44), "功能分解树", font=title, fill="#0F1F3D")
    root = (650, 120, 1150, 210)
    box(draw, root, "EduMind Agent\n智能学习协同系统", fill="#E8EEF5", outline="#2E74B5", font_size=29)
    groups = [
        ("身份与权限", ["演示登录", "角色识别", "Token 鉴权", "接口保护"]),
        ("学习管理", ["课程浏览", "目标创建", "任务拆解", "进度计算"]),
        ("AI 服务", ["课程问答", "计划生成", "笔记摘要", "Provider 切换"]),
        ("协作同步", ["协作消息", "活动日志", "SSE 推送", "多端刷新"]),
        ("项目支撑", ["JSON 持久化", "测试验证", "截图证据", "代码量统计"]),
    ]
    start_x = 80
    group_w = 310
    for i, (name, children) in enumerate(groups):
        gx = start_x + i * 340
        group_box = (gx, 330, gx + group_w, 410)
        arrow(draw, ((root[0] + root[2]) / 2, root[3]), (gx + group_w / 2, group_box[1]))
        box(draw, group_box, name, fill="#FFFFFF", font_size=26)
        for j, child in enumerate(children):
            y = 490 + j * 112
            child_box = (gx + 20, y, gx + group_w - 20, y + 72)
            arrow(draw, (gx + group_w / 2, group_box[3]), (gx + group_w / 2, child_box[1]))
            box(draw, child_box, child, fill="#FFFFFF", outline="#D8DEE9", font_size=23)
    img.save(path)


def save_use_case(path: Path):
    img = Image.new("RGB", (1800, 1120), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "UML 用例图（权限视角）", font=pil_font(42), fill="#0F1F3D")
    draw.rounded_rectangle((430, 130, 1500, 980), radius=20, outline="#2E74B5", width=4, fill="#F7FAFF")
    draw.text((470, 154), "系统边界：EduMind Agent", font=pil_font(30), fill="#1F4D78")
    actors = [("学生", 160, 260), ("教师", 160, 520), ("管理员", 160, 780)]
    for name, x, y in actors:
        draw.ellipse((x - 28, y - 80, x + 28, y - 24), outline="#52637A", width=3)
        draw.line((x, y - 24, x, y + 62), fill="#52637A", width=3)
        draw.line((x - 52, y + 8, x + 52, y + 8), fill="#52637A", width=3)
        draw.line((x, y + 62, x - 46, y + 128), fill="#52637A", width=3)
        draw.line((x, y + 62, x + 46, y + 128), fill="#52637A", width=3)
        draw.text((x - 42, y + 146), name, font=pil_font(28), fill="#172033")
    cases = [
        ("登录系统", 650, 270),
        ("管理目标/任务/笔记", 1030, 270),
        ("调用 AI 学习教练", 780, 470),
        ("查看课程与进度", 1200, 470),
        ("发送协作消息", 740, 690),
        ("查看活动日志", 1130, 690),
        ("维护课程资源", 950, 850),
    ]
    for label, cx, cy in cases:
        draw.ellipse((cx - 150, cy - 50, cx + 150, cy + 50), outline="#8EA7CB", fill="#FFFFFF", width=3)
        draw_centered_text(draw, (cx - 150, cy - 50, cx + 150, cy + 50), label, pil_font(24))
    links = [
        ((210, 280), (500, 270)), ((210, 280), (880, 270)), ((210, 280), (630, 470)), ((210, 280), (590, 690)),
        ((210, 540), (1050, 470)), ((210, 540), (980, 690)), ((210, 540), (800, 850)),
        ((210, 800), (500, 270)), ((210, 800), (800, 850)), ((210, 800), (980, 690)),
    ]
    for a, b in links:
        draw.line((a, b), fill="#667085", width=2)
    img.save(path)


def save_class_diagram(path: Path):
    img = Image.new("RGB", (1900, 1320), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "领域类图（核心业务对象）", font=pil_font(42), fill="#0F1F3D")
    classes = [
        ("User", ["id", "name", "role", "email"], ["canManageCourse()", "canReadLearningProfile()"], 80, 160),
        ("Course", ["id", "title", "teacherId", "tags"], ["toJSON()"], 520, 160),
        ("LearningGoal", ["ownerId", "courseId", "progress"], ["recalculateProgress()"], 980, 160),
        ("StudyTask", ["goalId", "status", "estimateMinutes"], ["start()", "complete()"], 1420, 160),
        ("LearningNote", ["ownerId", "courseId", "content"], ["summarize()"], 300, 620),
        ("AITutorService", ["provider", "templates"], ["ask()", "generatePlan()", "summarizeNote()"], 760, 620),
        ("PromptTemplate", ["system", "user"], ["render()"], 1220, 620),
        ("RoomMessage", ["roomId", "authorId", "content"], ["toJSON()"], 520, 980),
        ("ActivityLog", ["actorId", "type", "summary"], ["latest()"], 980, 980),
    ]
    positions = {}
    for name, attrs, methods, x, y in classes:
        w, h = 360, 250
        positions[name] = (x, y, x + w, y + h)
        draw.rounded_rectangle((x, y, x + w, y + h), radius=10, fill="#FFFFFF", outline="#8EA7CB", width=3)
        draw.rectangle((x, y, x + w, y + 56), fill="#E8EEF5", outline="#8EA7CB", width=2)
        draw_centered_text(draw, (x, y, x + w, y + 56), name, pil_font(26), fill="#0F1F3D")
        ty = y + 74
        for attr in attrs:
            draw.text((x + 22, ty), f"- {attr}", font=pil_font(21), fill="#172033")
            ty += 30
        draw.line((x, y + 156, x + w, y + 156), fill="#D8DEE9", width=2)
        ty = y + 174
        for method in methods:
            draw.text((x + 22, ty), f"+ {method}", font=pil_font(21), fill="#172033")
            ty += 30
    relations = [
        ("User", "Course", "教师维护"),
        ("User", "LearningGoal", "拥有"),
        ("LearningGoal", "StudyTask", "组合 1..*"),
        ("Course", "LearningGoal", "目标归属"),
        ("Course", "LearningNote", "笔记归属"),
        ("AITutorService", "PromptTemplate", "依赖"),
        ("AITutorService", "LearningGoal", "读取上下文"),
        ("User", "RoomMessage", "发送"),
        ("ActivityLog", "User", "记录行为"),
    ]
    for a, b, label in relations:
        ax1, ay1, ax2, ay2 = positions[a]
        bx1, by1, bx2, by2 = positions[b]
        start = ((ax1 + ax2) / 2, (ay1 + ay2) / 2)
        end = ((bx1 + bx2) / 2, (by1 + by2) / 2)
        arrow(draw, start, end, fill="#667085", width=2)
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        draw.rounded_rectangle((mx - 60, my - 18, mx + 60, my + 18), radius=8, fill="#FFFFFF", outline="#D8DEE9")
        draw_centered_text(draw, (mx - 60, my - 18, mx + 60, my + 18), label, pil_font(16))
    img.save(path)


def save_package_arch(path: Path):
    img = Image.new("RGB", (1900, 1180), "#F7F8FB")
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "包图、组件图与部署关系", font=pil_font(42), fill="#0F1F3D")
    packages = [
        ("client", "Web/PWA 前端\n视图、表单、状态、API Client", 70, 170, 430, 360),
        ("framework", "HTTP Server\nRouter / Error / SSE", 560, 170, 920, 360),
        ("application", "应用服务\nAuth / Learning / AI / Collaboration", 1050, 170, 1480, 360),
        ("domain", "领域对象\nUser / Goal / Task / Note / Message", 560, 520, 980, 720),
        ("infrastructure", "基础设施\nJsonDatabase / SeedData", 1100, 520, 1480, 720),
        ("LM Studio LLM", "OpenAI-compatible API\nqwen3.5-9b-glm5.1-distill-v1", 560, 850, 980, 1020),
        ("local data", "data/app-data.json\n手写知识库检索", 1100, 850, 1480, 1020),
    ]
    for title, desc, x1, y1, x2, y2 in packages:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#FFFFFF", outline="#B8C2D6", width=3)
        draw.rectangle((x1, y1, x2, y1 + 56), fill="#E8EEF5", outline="#B8C2D6", width=2)
        draw_centered_text(draw, (x1, y1, x2, y1 + 56), title, pil_font(25), fill="#0F1F3D")
        draw_centered_text(draw, (x1 + 20, y1 + 72, x2 - 20, y2 - 16), desc, pil_font(23), fill="#172033")
    arrows = [
        ((430, 265), (560, 265), "REST/SSE"),
        ((920, 265), (1050, 265), "Controller"),
        ((1265, 360), (770, 520), "调用领域模型"),
        ((1265, 360), (1290, 520), "读写仓储"),
        ((770, 720), (770, 850), "AI 上下文"),
        ((1290, 720), (1290, 850), "持久化"),
    ]
    for start, end, label in arrows:
        arrow(draw, start, end, fill="#52637A")
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        draw.text((mx - 58, my - 28), label, font=pil_font(18), fill="#52637A")
    draw.rounded_rectangle((1530, 220, 1820, 950), radius=22, outline="#2E74B5", width=4, fill="#FFFFFF")
    draw.text((1580, 260), "部署节点", font=pil_font(32), fill="#1F4D78")
    deploy = ["开发机/服务器", "Node.js 进程", "静态前端文件", "JSON 数据文件", "LM Studio: qwen3.5-9b"]
    for i, item in enumerate(deploy):
        y = 340 + i * 100
        box(draw, (1570, y, 1780, y + 64), item, fill="#F7FAFF", outline="#D8DEE9", font_size=20)
    img.save(path)


def save_sequence(path: Path):
    img = Image.new("RGB", (1900, 1000), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "顺序图：AI 学习计划生成", font=pil_font(42), fill="#0F1F3D")
    actors = [("学生浏览器", 180), ("ApiClient", 500), ("AITutorService", 850), ("LearningService", 1220), ("LLM Provider", 1580)]
    top, bottom = 170, 900
    for name, x in actors:
        box(draw, (x - 120, top, x + 120, top + 64), name, fill="#E8EEF5", outline="#8EA7CB", font_size=22)
        draw.line((x, top + 64, x, bottom), fill="#C7CEDA", width=3)
    steps = [
        (180, 500, 300, "选择目标并提交 /api/ai/plan"),
        (500, 850, 390, "携带 Bearer Token 调用后端"),
        (850, 1220, 480, "读取目标、课程、任务上下文"),
        (1220, 850, 570, "返回学习画像"),
        (850, 1580, 660, "渲染 Prompt 并调用 Provider"),
        (1580, 850, 750, "返回计划文本"),
        (850, 500, 820, "封装 AIResponse"),
        (500, 180, 880, "前端展示建议与后续操作"),
    ]
    for x1, x2, y, label in steps:
        arrow(draw, (x1, y), (x2, y), fill="#52637A")
        draw.text((min(x1, x2) + 24, y - 34), label, font=pil_font(19), fill="#172033")
    img.save(path)


def set_run_font(run, name=FONT_CJK, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_CJK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_CJK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    set_header_footer(section)


def set_header_footer(section):
    hp = section.header.paragraphs[0]
    hp.text = ""
    run = hp.add_run("EduMind Agent 智学伴 | 面向对象技术与方法结课设计")
    set_run_font(run, size=9, color=MUTED)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp = section.footer.paragraphs[0]
    fp.text = ""
    run = fp.add_run("项目源码：EduMindAgent | 文档日期：2026-05-23")
    set_run_font(run, size=9, color=MUTED)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_para(doc, text="", style=None, bold=False, color=None, size=None, align=None, after=None, before=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=LINE):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_FILL)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=INK, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=MUTED, italic=True)


def add_image(doc, path, caption, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def count_lines():
    total = 0
    include = {".js", ".mjs", ".css", ".html", ".json", ".sql", ".cmd"}
    ignore_dirs = {"node_modules", ".git", ".runtime", "data", "docs", "screenshots", "dist", "build", "coverage"}
    ignore_suffixes = (".generated.js", ".log", ".zip", ".docx", ".png")
    for path in ROOT.rglob("*"):
        if (
            path.is_file()
            and path.suffix in include
            and not any(part in ignore_dirs for part in path.parts)
            and not path.name.endswith(ignore_suffixes)
        ):
            total += sum(1 for line in path.read_text(encoding="utf-8").splitlines() if line.strip())
    return total


def cover(doc: Document, line_count: int):
    add_para(doc, "《面向对象技术与方法》结课设计", bold=True, size=16, color=MUTED, after=4)
    add_para(doc, "EduMind Agent 智学伴", bold=True, size=26, color=RGBColor(0, 0, 0), after=4)
    add_para(doc, "智能学习计划与协作系统项目文档", size=15, color=RGBColor(55, 55, 55), after=14)
    meta = [
        ("项目类型", "三人团队项目；Web 前端 + Node.js Server 端；接入 LM Studio 本地 LLM 服务"),
        ("项目边界", "课程学习目标、任务拆解、笔记沉淀、AI 辅导与协作同步"),
        ("代码规模", f"第一阶段可审查源码共 {line_count} 行；统计已排除 generated 文件、日志、截图和运行数据，后续通过真实模块扩展到课程要求规模"),
        ("运行地址", "默认 http://127.0.0.1:4077，截图验证端口 4081"),
        ("文档日期", "2026-05-23"),
    ]
    add_table(doc, ["项目元数据", "说明"], meta, [1900, 7460])
    add_para(
        doc,
        "本文档按照课程要求组织为《系统需求文档》《系统设计文档》《项目管理文档》三部分，并包含领域模型、UML 图、架构说明、运行截图、AI 应用情况和开发记录。",
        after=10,
    )
    doc.add_page_break()


def write_document():
    DOCS.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "feature": DOCS / "diagram-feature-tree.png",
        "usecase": DOCS / "diagram-use-case.png",
        "class": DOCS / "diagram-class.png",
        "package": DOCS / "diagram-package-architecture.png",
        "sequence": DOCS / "diagram-sequence-ai-plan.png",
    }
    save_feature_tree(diagrams["feature"])
    save_use_case(diagrams["usecase"])
    save_class_diagram(diagrams["class"])
    save_package_arch(diagrams["package"])
    save_sequence(diagrams["sequence"])

    line_count = count_lines()
    doc = Document()
    style_document(doc)
    cover(doc, line_count)

    add_para(doc, "目录", style="Heading 1")
    for item in [
        "一、系统需求文档",
        "二、系统设计文档",
        "三、项目管理文档",
        "附录：运行与提交说明",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_para(doc, "一、系统需求文档", style="Heading 1")
    add_para(doc, "1.1 项目可行性分析与领域边界", style="Heading 2")
    add_para(
        doc,
        "EduMind Agent 面向高校课程项目化学习场景，解决学生在多课程并行、任务拆解不清晰、复习材料分散、项目文档证据难以沉淀等方面的问题。系统领域边界限定为“学习计划与协作”，不覆盖教务选课、正式成绩录入、在线支付等校级系统能力。",
    )
    for text in [
        "服务用户：学生、任课教师和系统管理员。学生使用系统创建目标、拆解任务、记录笔记并调用 AI 学习教练；教师查看课程协作消息并提供指导；管理员负责系统配置与资源维护。",
        "核心价值：通过 AI 问答、计划生成和笔记摘要把学习活动变成可执行任务，并通过 SSE 同步保持多端信息一致。",
        "可行性：项目使用 Node.js 内置模块、浏览器原生 ES Modules 和 JSON 文件存储，无需外网安装依赖，便于课程提交、演示和复现。",
    ]:
        add_bullet(doc, text)
    add_para(doc, "竞品分析", style="Heading 3")
    add_table(
        doc,
        ["系统", "优势", "不足", "本项目差异"],
        [
            ("Notion/飞书文档", "文档与协作能力强", "学习计划、课程上下文和 AI 服务需要自行拼装", "将课程、目标、任务、笔记、AI 辅导放入同一领域模型"),
            ("豆包/通用 AI App", "对话体验好，模型能力强", "缺少课程项目管理与本地持久化", "AI 作为后端服务，与目标/任务/笔记微服务协同"),
            ("普通 Todo App", "任务管理轻量", "不理解课程场景和项目证据", "任务与课程、笔记、文档要求、协作消息关联"),
        ],
        [1800, 2500, 2500, 2560],
    )

    add_para(doc, "1.2 系统功能分析、原型设计及开发计划", style="Heading 2")
    add_image(doc, diagrams["feature"], "图 1 功能分解树：从身份、学习、AI、协作和项目支撑五类能力展开。", width=6.3)
    add_para(doc, "功能重要性与依赖关系", style="Heading 3")
    add_table(
        doc,
        ["功能", "权重", "依赖", "前端/后端边界"],
        [
            ("登录与鉴权", "10", "无", "前端保存 Token；后端签发和校验 Token"),
            ("课程/目标/任务管理", "10", "登录", "前端表单与列表；后端领域服务、仓储和进度计算"),
            ("AI 问答与计划生成", "9", "登录、学习上下文", "前端对话界面；后端 AITutorService 与 LLM Provider"),
            ("协作消息与活动日志", "8", "登录", "前端协作区；后端消息仓储、SSE 推送"),
            ("截图、测试、文档证据", "7", "项目可运行", "脚本生成截图；文档嵌入证据"),
        ],
        [2300, 900, 2100, 4060],
    )
    add_para(doc, "UML 用例图与权限分配", style="Heading 3")
    add_image(doc, diagrams["usecase"], "图 2 用例图：学生、教师、管理员三类参与者及可访问的系统能力。", width=6.3)
    add_para(doc, "客户端与服务端功能分工", style="Heading 3")
    add_table(
        doc,
        ["Client 端", "Server 端"],
        [
            ("登录页、总览页、学习页、AI 页、协作页；浏览器原生 ES Modules；localStorage 保存 Token。", "REST API、Token 鉴权、领域服务、JSON 持久化、AI Provider、SSE 事件推送。"),
            ("对目标、任务、笔记和消息进行表单交互，实时刷新界面状态。", "负责业务规则、权限边界、数据一致性和活动日志记录。"),
        ],
        [4680, 4680],
    )
    add_para(doc, "运行截图与 UI 说明", style="Heading 3")
    for image_name, caption in [
        ("screenshot-dashboard.png", "图 3 学习总览：聚合课程、目标进度、今日任务和 AI Provider 状态。"),
        ("screenshot-learning.png", "图 4 目标与任务：创建目标、拆解任务、保存学习笔记。"),
        ("screenshot-ai.png", "图 5 AI 学习教练：通过后端 AI 服务生成课程问答和学习计划。"),
        ("screenshot-team.png", "图 6 课程协作区：展示课程消息并通过服务端事件同步。"),
    ]:
        add_image(doc, DOCS / image_name, caption, width=6.3)

    add_para(doc, "开发计划安排", style="Heading 3")
    add_table(
        doc,
        ["日期", "迭代目标", "产出"],
        [
            ("2026-05-23", "完成需求解读、项目题目确定、前后端最小闭环", "项目目录、REST API、前端工作台、测试用例"),
            ("2026-05-24 至 2026-05-30", "完善领域模型、AI 知识库和文档图表", "领域类图、包图、AI 资源目录、截图"),
            ("2026-05-31 至 2026-06-08", "补充异常处理、权限策略、协作同步和更多测试", "接口测试、SSE 验证、边界用例"),
            ("2026-06-09 至 2026-06-16", "完成文档、部署说明和提交包整理", "DOCX 文档、README、代码量报告"),
            ("2026-06-17 至 2026-06-21", "最终演示、打包和乐学提交", "压缩包、演示视频或截图证据"),
        ],
        [1500, 3300, 4560],
    )

    add_para(doc, "1.3 领域模型描述", style="Heading 2")
    add_image(doc, diagrams["class"], "图 7 领域类图：User、Course、LearningGoal、StudyTask、LearningNote、AITutorService 等核心对象及关系。", width=6.3)
    add_image(doc, diagrams["package"], "图 8 包图/组件图/部署关系：client、framework、application、domain、infrastructure 与外部 LLM 协作。", width=6.3)
    add_image(doc, diagrams["sequence"], "图 9 顺序图：AI 学习计划生成时前端、服务、领域上下文和 LLM Provider 的协作。", width=6.3)
    add_para(
        doc,
        "领域对象采用 Entity + Repository 的基本结构。User 负责角色与权限判断；LearningGoal 根据 StudyTask 自动计算 progress；AITutorService 通过 PromptTemplate 渲染提示词，并调用 MockLLMProvider、LMStudioProvider 或 OpenAICompatibleProvider；ActivityLog 记录重要业务行为，供总览页展示。",
    )

    add_para(doc, "二、系统设计文档", style="Heading 1")
    add_para(doc, "2.1 技术选型", style="Heading 2")
    add_table(
        doc,
        ["技术项", "候选方案", "选择", "理由"],
        [
            ("后端语言", "Java / .NET / Node.js / PHP", "Node.js", "课程提交环境轻量；内置 HTTP、Crypto、Test Runner 可支撑无依赖项目"),
            ("前端", "Vue / React / Angular / 原生 Web", "原生 ES Modules", "不需要安装依赖，能直接服务静态文件并兼容移动端浏览器"),
            ("数据存储", "MySQL / SQLite / MongoDB / JSON", "JSON 文件", "团队演示阶段便于打包与复现；后续可替换为 SQLite/MySQL"),
            ("AI 接入", "固定云服务 / LM Studio / OpenAI 兼容 / 本地 Mock", "LM Studio Provider", "默认接入 qwen3.5-9b-glm5.1-distill-v1；Mock 仅用于离线回退和单元测试"),
            ("同步机制", "轮询 / WebSocket / SSE", "SSE", "服务端推送简单可靠，适合活动日志和协作消息同步"),
        ],
        [1500, 2200, 1500, 4160],
    )

    add_para(doc, "2.2 系统架构设计方案", style="Heading 2")
    add_para(
        doc,
        "系统采用前后端分离和领域分层结构。Server 端内部按 framework、application、domain、infrastructure 划分职责，虽然部署为一个 Node.js 进程，但每个 bounded context 都按照微服务边界设计，后续可以拆分为独立服务。",
    )
    for text in [
        "身份服务：负责登录、Token 签发、Token 校验和角色识别。",
        "学习服务：负责课程、目标、任务、笔记的业务规则和进度计算。",
        "AI 服务：负责提示词模板、本地课程知识库检索、Mock/LM Studio/OpenAI-compatible Provider 切换。",
        "协作服务：负责消息、活动日志和 SSE 事件广播。",
        "基础设施服务：负责 JSON 持久化、种子数据和脚本化截图验证。",
    ]:
        add_bullet(doc, text)
    add_para(doc, "跨服务协作与事务边界", style="Heading 3")
    add_para(
        doc,
        "学习服务在创建任务、完成任务后会同步更新 LearningGoal 的 progress，并通过 ActivityService 记录日志，再由 DomainEventBus 发布事件给 SyncHub。当前存储为单文件 JSON，跨边界事务由同一应用服务顺序保存；若拆成真实微服务，可演化为事件驱动的 Saga 模式。",
    )
    add_para(doc, "安全、并发、监控和部署", style="Heading 3")
    add_table(
        doc,
        ["主题", "方案"],
        [
            ("认证与鉴权", "JWT-lite HMAC Token；所有业务 API 通过 Authorization Bearer 校验；AI 服务调用必须登录。"),
            ("数据完整性", "应用服务集中校验必填字段、资源归属和状态流转；任务完成后统一回写目标进度。"),
            ("高并发应对", "团队项目以课程级并发为目标；可通过反向代理、进程集群、请求限流、缓存和数据库替换继续扩展。"),
            ("监控与日志", "ActivityLog 记录领域事件；运行时可扩展为结构化日志、健康检查和告警。"),
            ("部署方式", "Node.js 单进程部署静态前端和 API；LM Studio 在 172.25.160.1:1234 启动 OpenAI-compatible Server；自动部署可用 GitHub Actions 或服务器脚本。"),
            ("版本演化", "Provider、Repository、Service 分离，后续可替换数据库、前端框架或 LLM 供应商。"),
        ],
        [1900, 7460],
    )

    add_para(doc, "2.3 技术关键点", style="Heading 2")
    add_para(doc, "关键点一：可替换 LLM Provider", style="Heading 3")
    add_para(
        doc,
        "AITutorService 不直接绑定某个云厂商，而是依赖 Provider 的 complete(messages) 能力。团队部署默认使用 LMStudioProvider 指向 http://172.25.160.1:1234/v1/chat/completions，并调用 qwen3.5-9b-glm5.1-distill-v1；MockLLMProvider 仅用于离线回退和单元测试。项目提供 npm run test:lmstudio 脚本用于提交前真实请求本地模型。",
    )
    add_para(doc, "关键点二：领域对象驱动的任务进度计算", style="Heading 3")
    add_para(
        doc,
        "LearningGoal 暴露 recalculateProgress(tasks) 方法，StudyTask 暴露 start() 与 complete() 方法。应用服务只编排对象协作，进度算法留在领域对象内部，体现面向对象封装。",
    )
    add_para(doc, "关键点三：SSE 多端协同", style="Heading 3")
    add_para(
        doc,
        "DomainEventBus 发布 goal.changed、task.changed、note.changed、message.created、activity.created 等事件；SyncHub 将事件转换为 text/event-stream。多个浏览器窗口登录后能共享和互通信息。",
    )
    add_para(doc, "关键点四：无依赖可运行工程", style="Heading 3")
    add_para(
        doc,
        "项目不要求 npm install，后端使用 Node.js 内置 http、crypto、fs、test 等模块，前端使用浏览器原生能力。这样可以降低课程提交后的复现风险。",
    )

    add_para(doc, "2.4 系统附加说明", style="Heading 2")
    add_table(
        doc,
        ["项目", "说明"],
        [
            ("开发环境", "Windows；Node.js 20+；浏览器 Chrome/Edge；LM Studio 本地模型服务 qwen3.5-9b-glm5.1-distill-v1。"),
            ("运行方法", "进入 EduMindAgent，执行 npm start，浏览器访问 http://127.0.0.1:4077。"),
            ("测试方法", "执行 npm test 验证核心流程；执行 npm run test:lmstudio 验证真实 LM Studio 模型调用。"),
            ("代码规模", f"scripts/linecount.mjs 当前统计可审查源码 {line_count} 行；统计口径排除 generated 文件、日志、截图、文档和运行数据。"),
            ("安装/卸载", "安装：解压项目并安装 Node.js；卸载：删除项目目录即可，数据位于 data/app-data.json。"),
            ("可执行程序", "本项目为 Web 系统，未额外打包桌面可执行程序；可选后续使用 pkg/electron 封装。"),
        ],
        [1900, 7460],
    )

    add_para(doc, "三、项目管理文档", style="Heading 1")
    add_para(doc, "3.1 项目组成员及分工", style="Heading 2")
    add_table(
        doc,
        ["成员", "角色", "承担工作", "贡献占比"],
        [
            ("顾昌宇", "组长/架构与集成", "需求拆解、领域模型、架构设计、Git 仓库管理、最终集成与提交包审核", "36%"),
            ("王宗翰", "后端与 AI 服务", "REST API、鉴权、领域服务、LM Studio Provider 接入、本地课程知识库与 AI 测试", "34%"),
            ("张宇航", "前端、测试与文档", "Web/PWA 工作台、响应式 UI、单元测试、截图验证、结课设计文档整理", "30%"),
        ],
        [1500, 1900, 4460, 1500],
    )

    add_para(doc, "3.2 AI 应用情况报告", style="Heading 2")
    for text in [
        "需求理解：使用 AI 辅助抽取结课设计文档中的交付要求、评分关注点和文档结构。",
        "代码生成：使用 AI 生成 Node.js 后端、Web 前端、测试脚本和截图脚本，并人工通过运行测试与截图验证结果。",
        "AI 服务设计：系统自身内置 AITutorService、PromptTemplate、MockLLMProvider、LMStudioProvider 和 OpenAICompatibleProvider。团队部署时由 LM Studio 提供 qwen3.5-9b-glm5.1-distill-v1 模型服务，地址为 http://172.25.160.1:1234；提交前通过 npm run test:lmstudio 验证真实模型调用。",
        "文档生成：使用 AI 协助组织需求文档、设计文档和项目管理文档；项目总结章节按课程要求留给学生自行补写。",
    ]:
        add_bullet(doc, text)

    add_para(doc, "3.3 实际开发记录", style="Heading 2")
    add_table(
        doc,
        ["日期", "开发活动", "相关产出"],
        [
            ("2026-05-23 11:27", "读取结课设计要求，确定智能学习协同系统主题", "需求拆解、项目题目"),
            ("2026-05-23 11:36", "创建项目目录与后端基础框架", "framework、infrastructure、domain、application 目录"),
            ("2026-05-23 11:45", "实现领域对象、应用服务和 REST API", "User、LearningGoal、StudyTask、AITutorService 等类"),
            ("2026-05-23 12:05", "实现浏览器前端工作台", "登录、总览、学习、AI、协作页面"),
            ("2026-05-23 12:20", "补充测试与 AI 课程知识库", "Node.js 测试、AI 本地知识库初版"),
            ("2026-05-23 12:30", "接入 LM Studio Provider 接口并清理代码统计口径", "LMStudioProvider、手写课程知识库、有效代码统计脚本"),
            ("2026-05-23 17:40", "配置团队本地大模型服务", "默认接入 qwen3.5-9b-glm5.1-distill-v1，新增 npm run test:lmstudio 验证脚本"),
            ("2026-05-23 12:35", "使用 Playwright 进行页面验证并生成截图", "dashboard、learning、ai、team 四张截图"),
            ("2026-05-23 12:50", "编写结课设计文档与 UML/架构图", "DOCX 文档、功能树、用例图、类图、包图、顺序图"),
        ],
        [1600, 3000, 4760],
    )
    add_para(
        doc,
        "Commit 记录建议：正式提交到 Git 仓库时，可按“初始化项目”“实现后端服务”“实现前端工作台”“接入 LM Studio Provider”“配置 qwen3.5-9b 本地模型”“扩展真实业务模块”“补充测试与文档”组织提交历史。",
    )
    add_para(doc, "3.4 项目总结", style="Heading 2")
    add_para(
        doc,
        "课程要求说明“项目总结不要使用 AI 生成”。因此本节仅保留位置，建议学生在最终提交前根据真实开发体验手写 300-500 字反思，重点写遇到的问题、解决过程、协作体验和对面向对象方法的理解。",
    )

    add_para(doc, "附录：运行与提交说明", style="Heading 1")
    add_para(doc, "提交包建议包含以下内容：", style="Heading 2")
    for text in [
        "EduMindAgent/ 项目源码目录。",
        "EduMindAgent/docs/EduMindAgent_结课设计文档.docx。",
        "运行截图 PNG 文件，用于无法现场运行时的证据补充。",
        "README.md、package.json、测试脚本和代码量统计脚本。",
    ]:
        add_bullet(doc, text)
    add_para(
        doc,
        "乐学提交截止时间为 2026-06-21 23:59。建议在提交前重新执行 npm test、npm run linecount，并打开本地页面确认截图所示功能仍可运行。",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    write_document()
