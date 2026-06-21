from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
PROTOTYPES = ROOT / "prototypes"
REAL_SCREENSHOTS = DOCS / "real-screenshots"

FONT_CJK = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(23, 32, 51)
MUTED = RGBColor(102, 112, 133)
HEADER_FILL = "F2F4F7"
SOFT_FILL = "E8EEF5"
LINE = "D8DEE9"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

PROJECT_NAME = "EduMind Agent 智学伴"
TEAM = {
    "leader": "王宗翰",
    "members": ["顾昌宇", "张宇航"],
}


def font_path() -> str:
    for candidate in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if Path(candidate).exists():
            return candidate
    return ""


def pil_font(size: int):
    path = font_path()
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def set_run_font(run, name=FONT_CJK, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def style_document(doc: Document, title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CJK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_CJK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run(f"{PROJECT_NAME} | {title}")
    set_run_font(run, size=9, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    run = footer.add_run("面向对象技术与方法结课项目文档")
    set_run_font(run, size=9, color=MUTED)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_para(doc, text="", style=None, bold=False, color=None, size=None, align=None, before=0, after=6):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold)
    if align is not None:
        paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=before, after=after)
    return paragraph


def add_title_page(doc, title: str, subtitle: str):
    add_para(doc, "《面向对象技术与方法》课程结课项目", size=13, color=MUTED, after=4)
    add_para(doc, PROJECT_NAME, bold=True, size=24, color=INK, after=4)
    add_para(doc, title, bold=True, size=18, color=BLUE, after=14)
    add_para(doc, subtitle, size=12, color=INK, after=18)
    add_table(
        doc,
        ["项目项", "内容"],
        [
            ("项目定位", "面向课程学习、任务拆解、AI 辅导和协作记录的 Web 系统"),
            ("技术栈", "浏览器原生前端、Node.js 后端、JSON 文件持久化、LM Studio/OpenAI-compatible AI 接入"),
            ("小组成员", f"组长：{TEAM['leader']}；组员：{'、'.join(TEAM['members'])}"),
            ("文档说明", "根据 README、源码结构、开发文档、测试记录、截图和 UML 图整理，不虚构未实现功能。"),
        ],
        [1700, 7660],
    )
    doc.add_page_break()


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, after=0, line_spacing=1.10)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_borders(table, color=LINE):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], HEADER_FILL)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=INK, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(str(value)) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=10.2, align=align)
    add_para(doc, "", after=2)
    return table


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, after=4, line_spacing=1.167)
        run = paragraph.add_run(item)
        set_run_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(paragraph, after=4, line_spacing=1.167)
        run = paragraph.add_run(item)
        set_run_font(run, size=11)


def add_caption(doc, text):
    paragraph = add_para(doc, text, size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8)
    paragraph.paragraph_format.keep_with_next = False


def add_image(doc, image_path: Path, caption: str, width=6.1):
    if not image_path.exists():
        add_para(doc, f"图片缺失：{caption}", italic=True, color=MUTED)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def wrap_text(draw, text, font, max_width):
    lines = []
    current = ""
    for char in text:
        trial = current + char
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_box(draw, xy, text, fill="#FFFFFF", outline="#9AA8BC", font_size=26):
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=3)
    font = pil_font(font_size)
    x1, y1, x2, y2 = xy
    lines = wrap_text(draw, text, font, int(x2 - x1 - 28))
    line_h = font_size + 6
    y = y1 + ((y2 - y1) - line_h * len(lines)) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font, fill="#172033")
        y += line_h


def draw_arrow(draw, start, end):
    draw.line([start, end], fill="#52637A", width=3)
    ex, ey = end
    sx, _ = start
    direction = 1 if ex >= sx else -1
    draw.polygon([(ex, ey), (ex - 14 * direction, ey - 7), (ex - 14 * direction, ey + 7)], fill="#52637A")


def make_feature_tree(path: Path):
    img = Image.new("RGB", (1800, 1120), "#F7F8FB")
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "功能分解树", font=pil_font(42), fill="#0F1F3D")
    root = (650, 130, 1150, 220)
    draw_box(draw, root, "EduMind Agent\n智能学习协作系统", fill="#E8EEF5", outline="#2E74B5", font_size=28)
    groups = [
        ("身份与权限", ["演示登录", "角色识别", "Token 鉴权", "成员管理"]),
        ("学习管理", ["课程浏览", "目标创建", "任务拆解", "笔记沉淀"]),
        ("作业评测", ["作业发布", "提交记录", "Rubric 评分", "错题练习"]),
        ("AI 辅助", ["学习计划", "问答摘要", "自检建议", "教师干预草稿"]),
        ("协作统计", ["协作消息", "活动日志", "SSE 同步", "统计报表"]),
    ]
    for i, (name, children) in enumerate(groups):
        gx = 70 + i * 345
        group_box = (gx, 340, gx + 300, 420)
        draw_arrow(draw, ((root[0] + root[2]) / 2, root[3]), (gx + 150, group_box[1]))
        draw_box(draw, group_box, name, font_size=25)
        for j, child in enumerate(children):
            child_box = (gx + 25, 500 + j * 112, gx + 275, 572 + j * 112)
            draw_arrow(draw, (gx + 150, group_box[3]), (gx + 150, child_box[1]))
            draw_box(draw, child_box, child, outline="#D8DEE9", font_size=22)
    img.save(path)


def make_use_case(path: Path):
    img = Image.new("RGB", (1800, 1120), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "UML 用例图（角色视角）", font=pil_font(42), fill="#0F1F3D")
    draw.rounded_rectangle((420, 130, 1510, 980), radius=20, outline="#2E74B5", width=4, fill="#F7FAFF")
    draw.text((460, 158), "系统边界：EduMind Agent", font=pil_font(30), fill="#1F4D78")
    actors = [("学生", 170, 270), ("教师", 170, 540), ("管理员", 170, 790)]
    for name, x, y in actors:
        draw.ellipse((x - 28, y - 80, x + 28, y - 24), outline="#52637A", width=3)
        draw.line((x, y - 24, x, y + 62), fill="#52637A", width=3)
        draw.line((x - 52, y + 8, x + 52, y + 8), fill="#52637A", width=3)
        draw.line((x, y + 62, x - 44, y + 125), fill="#52637A", width=3)
        draw.line((x, y + 62, x + 44, y + 125), fill="#52637A", width=3)
        draw.text((x - 42, y + 142), name, font=pil_font(28), fill="#172033")
    cases = [
        ("登录系统", 630, 260),
        ("管理学习目标/任务", 1040, 260),
        ("调用 AI 学习助手", 790, 455),
        ("提交作业/练习", 1220, 455),
        ("查看学生证据", 760, 665),
        ("生成干预建议", 1170, 665),
        ("管理班级与权限", 940, 850),
    ]
    for label, cx, cy in cases:
        draw.ellipse((cx - 155, cy - 48, cx + 155, cy + 48), outline="#8EA7CB", fill="#FFFFFF", width=3)
        draw_box(draw, (cx - 155, cy - 48, cx + 155, cy + 48), label, outline="#FFFFFF", font_size=23)
    links = [
        ((220, 285), (475, 260)), ((220, 285), (885, 260)), ((220, 285), (635, 455)), ((220, 285), (1065, 455)),
        ((220, 555), (605, 665)), ((220, 555), (1015, 665)), ((220, 555), (785, 850)),
        ((220, 805), (475, 260)), ((220, 805), (785, 850)),
    ]
    for start, end in links:
        draw.line((start, end), fill="#667085", width=2)
    img.save(path)


def make_class_diagram(path: Path):
    img = Image.new("RGB", (1900, 1320), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "领域类图（核心对象）", font=pil_font(42), fill="#0F1F3D")
    classes = [
        ("User", "name\nrole\nemail\nstatus", 90, 170),
        ("Course", "title\nteacherId\ntags", 520, 170),
        ("LearningGoal", "ownerId\ncourseId\nprogress\nrecalculateProgress()", 960, 170),
        ("StudyTask", "goalId\nstatus\ndueDate\ncomplete()", 1400, 170),
        ("LearningNote", "ownerId\ncourseId\ncontent\ntags", 300, 620),
        ("AITutorService", "ask()\ngeneratePlan()\nsummarizeNote()\nreviewSubmission()", 760, 620),
        ("Assignment", "courseId\nrubricId\ndueAt\nstatus", 1220, 620),
        ("Question/Practice", "questionIds\nanswerRecords\nmasteryScore", 300, 980),
        ("RoomMessage", "roomId\nauthorId\ncontent", 760, 980),
        ("ActivityLog", "actorId\ntype\nsummary\noccurredAt", 1220, 980),
    ]
    positions = {}
    for name, body, x, y in classes:
        positions[name] = (x, y, x + 360, y + 220)
        draw.rounded_rectangle((x, y, x + 360, y + 220), radius=10, fill="#FFFFFF", outline="#8EA7CB", width=3)
        draw.rectangle((x, y, x + 360, y + 54), fill="#E8EEF5", outline="#8EA7CB", width=2)
        draw_box(draw, (x, y, x + 360, y + 54), name, outline="#E8EEF5", fill="#E8EEF5", font_size=25)
        for idx, line in enumerate(body.splitlines()):
            draw.text((x + 22, y + 78 + idx * 31), f"- {line}", font=pil_font(21), fill="#172033")
    relations = [
        ("User", "LearningGoal"), ("LearningGoal", "StudyTask"), ("Course", "Assignment"),
        ("Assignment", "Question/Practice"), ("AITutorService", "LearningGoal"),
        ("User", "RoomMessage"), ("ActivityLog", "User"), ("Course", "LearningNote"),
    ]
    for left, right in relations:
        a = positions[left]
        b = positions[right]
        draw_arrow(draw, ((a[0] + a[2]) / 2, (a[1] + a[3]) / 2), ((b[0] + b[2]) / 2, (b[1] + b[3]) / 2))
    img.save(path)


def make_package_diagram(path: Path):
    img = Image.new("RGB", (1900, 1180), "#F7F8FB")
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "包图与逻辑服务划分", font=pil_font(42), fill="#0F1F3D")
    boxes = [
        ("client", "原生 Web/PWA\n视图、状态、API Client", 80, 170),
        ("gateway-service", "静态前端托管\n鉴权、转发、聚合", 540, 170),
        ("identity-service", "用户、角色、班级\nToken 校验", 1000, 170),
        ("learning-service", "课程、目标、任务\n笔记、学习上下文", 1460, 170),
        ("assessment-service", "作业、题库、练习\n评分、错题本", 310, 540),
        ("ai-service", "Prompt、Provider\n学生/教师 AI 工作流", 770, 540),
        ("collaboration-service", "消息、活动日志\n事件与 SSE", 1230, 540),
        ("analytics/report", "统计看板\n学习报告", 770, 890),
    ]
    for title, desc, x, y in boxes:
        draw.rounded_rectangle((x, y, x + 360, y + 160), radius=18, fill="#FFFFFF", outline="#B8C2D6", width=3)
        draw.rectangle((x, y, x + 360, y + 52), fill="#E8EEF5", outline="#B8C2D6", width=2)
        draw_box(draw, (x, y, x + 360, y + 52), title, fill="#E8EEF5", outline="#E8EEF5", font_size=24)
        draw_box(draw, (x + 20, y + 64, x + 340, y + 145), desc, outline="#FFFFFF", font_size=22)
    arrows = [
        ((440, 250), (540, 250)), ((900, 250), (1000, 250)), ((1360, 250), (1460, 250)),
        ((720, 330), (490, 540)), ((720, 330), (950, 540)), ((900, 330), (1410, 540)),
        ((950, 700), (950, 890)), ((1410, 700), (1130, 890)),
    ]
    for start, end in arrows:
        draw_arrow(draw, start, end)
    img.save(path)


def make_sequence_diagram(path: Path):
    img = Image.new("RGB", (1900, 1000), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "动态 UML：AI 学习计划生成顺序图", font=pil_font(42), fill="#0F1F3D")
    actors = [("学生浏览器", 180), ("Gateway", 500), ("AI Service", 850), ("Learning Service", 1230), ("LM Studio/Mock", 1600)]
    for name, x in actors:
        draw_box(draw, (x - 125, 160, x + 125, 225), name, fill="#E8EEF5", outline="#8EA7CB", font_size=22)
        draw.line((x, 225, x, 910), fill="#C7CEDA", width=3)
    steps = [
        (180, 500, 310, "POST /api/student-ai/daily-plan"),
        (500, 850, 400, "转发用户上下文"),
        (850, 1230, 490, "读取目标、任务、笔记上下文"),
        (1230, 850, 580, "返回学习画像"),
        (850, 1600, 670, "渲染 Prompt 并调用 Provider"),
        (1600, 850, 760, "返回结构化结果或 fallback"),
        (850, 500, 835, "保存 AI result"),
        (500, 180, 890, "返回计划卡片"),
    ]
    for x1, x2, y, label in steps:
        draw_arrow(draw, (x1, y), (x2, y))
        draw.text((min(x1, x2) + 24, y - 32), label, font=pil_font(18), fill="#172033")
    img.save(path)


def make_diagrams(tmpdir: Path):
    diagrams = {
        "feature": tmpdir / "feature.png",
        "usecase": tmpdir / "usecase.png",
        "class": tmpdir / "class.png",
        "package": tmpdir / "package.png",
        "sequence": tmpdir / "sequence.png",
    }
    make_feature_tree(diagrams["feature"])
    make_use_case(diagrams["usecase"])
    make_class_diagram(diagrams["class"])
    make_package_diagram(diagrams["package"])
    make_sequence_diagram(diagrams["sequence"])
    return diagrams


def find_screenshots():
    candidates = [
        ("学生 AI 首页", PROTOTYPES / "student-ai-first-v2" / "screenshots" / "desktop-ai.png"),
        ("学习任务页面", PROTOTYPES / "student-ai-first-v2" / "screenshots" / "desktop-learn.png"),
        ("作业详情页面", PROTOTYPES / "student-ai-first-v2" / "screenshots" / "desktop-assignment-preview.png"),
        ("教师 AI 工作台", PROTOTYPES / "teacher-ai-first" / "screenshots" / "desktop-ai.png"),
        ("教师批改页面", PROTOTYPES / "teacher-ai-first" / "screenshots" / "desktop-grading-detail.png"),
        ("协作页面", PROTOTYPES / "teacher-ai-first" / "screenshots" / "desktop-collaboration.png"),
    ]
    return [(label, path) for label, path in candidates if path.exists()]


def find_design_screenshots():
    return {
        "overview": REAL_SCREENSHOTS / "live-login-overview.png",
        "student_ai": REAL_SCREENSHOTS / "live-student-ai.png",
        "teacher_ai": REAL_SCREENSHOTS / "live-teacher-ai.png",
        "student_learn": REAL_SCREENSHOTS / "live-student-learning.png",
        "teacher_assignments": REAL_SCREENSHOTS / "live-teacher-assignments.png",
        "student_ai_insight": REAL_SCREENSHOTS / "live-student-ai-insight.png",
        "student_note_ai": REAL_SCREENSHOTS / "live-student-note-ai-result.png",
        "teacher_ai_plan": REAL_SCREENSHOTS / "live-teacher-ai-plan.png",
        "teacher_intervention": REAL_SCREENSHOTS / "live-teacher-intervention.png",
        "teacher_collaboration": REAL_SCREENSHOTS / "live-collaboration.png",
        "teacher_collaboration_detail": REAL_SCREENSHOTS / "live-collaboration-detail.png",
        "student_assignment_preview": REAL_SCREENSHOTS / "live-student-assignment-preview.png",
        "teacher_grading": REAL_SCREENSHOTS / "live-teacher-grading.png",
        "teacher_analytics": REAL_SCREENSHOTS / "live-report.png",
    }


def build_requirement_doc(diagrams):
    doc = Document()
    title = "系统需求文档"
    style_document(doc, title)
    add_title_page(doc, title, "从课程作业要求、README、源码结构、截图和 UML 图整理系统需求。")

    add_para(doc, "一、项目可行性分析", style="Heading 1")
    add_para(doc, "本项目适合作为课程结课项目实现。系统选题来自真实学习场景：学生需要把课程目标拆成任务，教师需要看到学生学习过程证据，小组项目还需要协作记录和文档支撑。实现上采用浏览器原生前端、Node.js 后端和 JSON 文件存储，部署门槛低，适合课堂演示和提交复现。")
    add_table(doc, ["分析维度", "结论", "依据"], [
        ("技术可行性", "可行。项目使用 Node.js 20+、浏览器原生 ES Modules、JSON 文件和 Server-Sent Events，不依赖复杂外部组件。", "README、package.json、client、services、shared 目录。"),
        ("业务可行性", "可行。学习目标、任务、笔记、作业、题库、练习、AI 辅导和协作消息都能对应课程学习过程。", "learning-service、assessment-service、ai-service、collaboration-service。"),
        ("经济可行性", "可行。课程项目阶段不需要购买云服务器或数据库，AI 可使用本地 LM Studio 或 Mock Provider。", "README 的 LM Studio 配置和 npm run start:mock。"),
        ("操作可行性", "可行。系统通过浏览器访问，学生和教师使用常见的学习、作业、消息和 AI 面板交互。", "client/src/views、prototypes 截图。"),
        ("进度可行性", "可行。系统已经通过自动测试，文档和截图能支撑结课提交。", "npm test、npm run test:services、npm run test:ai。"),
    ], [1700, 3800, 3860])
    add_para(doc, "项目风险也比较明确。当前系统不是生产级学校平台，不处理正式选课、学籍、支付、考试监控和完整 OAuth 登录。密码注册、session 失效、多租户和数据库迁移在开发文档中属于后续目标，本次需求文档只把它们写成改进方向。")

    add_para(doc, "二、领域边界与目标用户", style="Heading 1")
    add_para(doc, "系统领域边界限定在“课程学习支持、AI 辅导、作业评测和协作记录”。它解决的是课程项目学习中的执行、反馈和证据问题，不替代学校教务系统，也不承担正式成绩管理职责。")
    add_table(doc, ["边界内", "说明"], [
        ("课程学习管理", "课程、学习目标、学习任务、学习笔记、学习路径和任务进度。"),
        ("作业与练习", "作业发布、学生提交、Rubric 评分、题库、练习会话、错题本和掌握度。"),
        ("AI 辅助", "学生计划、薄弱点分析、任务草稿、作业指导、提交自检、笔记整理，以及教师讲评、反馈和干预草稿。"),
        ("协作与证据", "房间消息、回复、提及、活动日志、协作任务、决策、资源和 SSE 事件。"),
        ("统计与报告", "学习概览、课程统计、学生画像、报告摘要和 AI 使用记录。"),
    ], [2300, 7060])
    add_table(doc, ["边界外", "不纳入本次课程项目的原因"], [
        ("正式教务选课", "需要对接学校教务系统和真实组织数据，超出课程项目范围。"),
        ("正式成绩归档", "本项目可保存练习和评分建议，但不作为学校成绩系统。"),
        ("在线考试监管", "摄像头监考、防作弊、考试锁屏等能力不属于学习计划系统。"),
        ("支付和商业运营", "项目不涉及收费、订单和发票。"),
        ("生产级账号体系", "当前为轻量登录；密码注册、session 管理、多租户作为后续规划。"),
    ], [2300, 7060])
    add_para(doc, "目标用户分为学生、教师和管理员/教师管理员。三类用户的核心目标不同，因此页面入口、权限和 AI 功能也不同。")
    add_table(doc, ["用户", "典型场景", "主要诉求", "系统响应"], [
        ("学生", "上课后整理笔记、拆分课程作业、提交前自检、根据错题复习。", "希望任务清楚、反馈及时、AI 建议能直接转成行动。", "学生 AI 工作区、学习目标、任务草稿、作业指导、提交自检、错题练习。"),
        ("教师", "查看学生作业、了解学习过程、批改反馈、安排干预。", "希望看见过程证据，而不是只看到最终提交。", "教师 AI 工作台、学生画像、作业批改、学生 AI 时间线、干预草稿。"),
        ("管理员/教师管理员", "维护班级、用户、分组和基础权限。", "希望系统演示数据清楚，角色入口可控。", "班级成员管理、角色权限查看、运维和统计入口。"),
    ], [1300, 2850, 2450, 2760])

    add_para(doc, "三、用户需求", style="Heading 1")
    add_para(doc, "用户需求分为功能需求和非功能需求。功能需求对应实际模块，非功能需求约束可用性、可复现性、安全边界和课程提交质量。")
    add_para(doc, "3.1 功能需求", style="Heading 2")
    add_table(doc, ["编号", "优先级", "需求", "验收方式"], [
        ("R1", "高", "用户可以登录系统并按 student、teacher、admin 角色进入不同页面。", "使用演示账号登录，检查学生端、教师端和管理入口差异。"),
        ("R2", "高", "学生可以查看课程、创建目标、拆分任务、记录笔记，并根据任务完成情况更新目标进度。", "创建目标和任务，完成任务后检查 LearningGoal progress。"),
        ("R3", "高", "学生可以使用 AI 生成每日计划、薄弱点分析、任务草稿、作业指导、提交自检和笔记整理。", "调用 /api/student-ai/*，检查 result、actions 和 fallback。"),
        ("R4", "高", "教师可以查看学生画像、作业提交、AI 时间线和过程证据，并生成干预或反馈草稿。", "打开教师端学生详情和批改页面，检查 teacher-ai 结果和草稿。"),
        ("R5", "高", "系统支持作业、提交、Rubric、评分、题库、练习、错题本和掌握度记录。", "运行 assessment-service 和 gateway 相关测试。"),
        ("R6", "中", "系统支持协作房间、消息、回复、提及、协作任务、资源、决策和交接记录。", "通过 collaboration-service 接口和协作页面检查。"),
        ("R7", "中", "系统支持活动日志和 SSE 事件推送，页面可以感知任务、消息或活动变化。", "连接 /api/events，发送内部事件后检查 ready/event 数据。"),
        ("R8", "中", "系统支持统计、报告和运维导入审计等扩展能力。", "检查 analytics、report、operations 服务接口和测试结果。"),
    ], [700, 900, 5200, 2560])
    add_para(doc, "3.2 非功能需求", style="Heading 2")
    add_table(doc, ["类别", "需求说明", "当前实现依据"], [
        ("可运行性", "项目应能在本地 Windows + Node.js 环境运行，尽量减少安装依赖。", "README、start-local.cmd、npm start、npm run start:services:mock。"),
        ("可测试性", "关键业务流程需要有自动测试，提交前能复现通过结果。", "test/*.mjs，npm test、npm run test:services、npm run test:ai。"),
        ("可维护性", "业务代码按服务、application、domain、infrastructure、client views/widgets/state 拆分。", "services、shared、client/src 目录结构。"),
        ("可解释性", "AI 输出需要保留 provider、generatedAt、rawText 或 fallback 信息，便于解释来源。", "AI result、ProviderHealth、AI workflow normalize。"),
        ("安全边界", "普通用户不能访问内部接口，学生不能读取他人 AI 结果，教师功能需要 teacher/admin 角色。", "userContext、requireInternal、ai-route-guards.test.mjs。"),
        ("数据可迁移", "课程阶段使用 JSON 文件，领域对象和 Repository 抽象保留未来迁移数据库的可能。", "shared/data/repository.js、各服务 domain repository。"),
    ], [1600, 4400, 3360])
    add_para(doc, "3.3 典型用户故事", style="Heading 2")
    add_table(doc, ["角色", "用户故事", "完成标准"], [
        ("学生", "作为学生，我希望 AI 根据今天的课程和作业给出学习计划，这样我可以直接安排学习任务。", "生成 daily_plan 后能看到摘要、行动项和建议问题。"),
        ("学生", "作为学生，我希望提交作业前先做 AI 自检，这样可以发现明显遗漏。", "submission_check 返回完成度、问题、优点和修改建议。"),
        ("教师", "作为教师，我希望查看学生 AI 学习时间线，这样可以判断学生是否真正有过程投入。", "教师端能查看指定学生 results 和 timeline。"),
        ("教师", "作为教师，我希望根据学生证据生成干预草稿，但最终由我确认是否发送。", "teacher_ai_draft 处于 draft，确认后状态变为 sent 或 saved。"),
        ("管理员", "作为管理员，我希望能查看班级和成员基础信息，这样演示数据更容易组织。", "identity-admin 页面和 identity-service 管理接口可用。"),
    ], [1200, 5300, 2860])

    add_para(doc, "四、竞品分析", style="Heading 1")
    add_para(doc, "竞品分析用于明确项目定位。本项目没有试图做成全能学习平台，而是在课程项目范围内把学习任务、作业证据、AI 辅助和教师反馈放到同一条流程中。")
    add_table(doc, ["参照系统", "优势", "不足", "本项目取舍"], [
        ("Notion/飞书文档", "文档、表格和协作能力成熟，适合自由记录。", "课程目标、作业提交、AI 自检和教师端证据查看需要用户手工组织。", "保留笔记和协作思路，但把课程、任务、作业和 AI 结果做成结构化对象。"),
        ("通用 AI 对话工具", "模型能力强，问答体验好，适合临时咨询。", "不理解本系统内的课程、任务、错题、提交和教师反馈状态。", "AI 作为后端服务读取学习上下文，输出可保存的 result、draft 和 action。"),
        ("普通 Todo/学习打卡应用", "任务管理简单，适合个人计划。", "缺少课程知识、作业评测、错题掌握度和教师反馈。", "任务与课程目标、作业、笔记和 AI 草稿关联。"),
        ("在线教学平台", "课程、作业和班级管理完整。", "一般偏教学管理，不强调 AI 过程证据和学生自我规划。", "不做正式教务，只聚焦课程项目学习闭环。"),
        ("题库/刷题系统", "练习和判题能力强。", "难以覆盖学习计划、笔记、教师干预和项目协作。", "题库和练习作为评测子模块，与学习目标和 AI 建议联动。"),
    ], [1600, 2500, 2600, 2660])
    add_para(doc, "与这些参照系统相比，本项目的特色是“过程证据”。学生不是只获得 AI 回答，教师也不是只看最终作业，而是能看到 AI 结果、行动项、任务草稿、提交自检、练习和协作记录。")

    add_para(doc, "五、功能分解树", style="Heading 1")
    add_image(doc, diagrams["feature"], "图 1 功能分解树：系统围绕身份、学习、作业评测、AI 辅助和协作统计组织功能。")
    add_para(doc, "功能分解树体现了项目的实现重点。身份与权限是所有业务入口；学习管理和作业评测是课程项目的主体；AI 辅助负责计划、问答、自检和教师草稿；协作统计用于保留过程证据。")
    add_table(doc, ["一级功能", "二级功能", "说明"], [
        ("身份与权限", "演示登录、角色识别、Token 鉴权、成员管理", "保证学生、教师、管理员进入对应功能区域。"),
        ("学习管理", "课程、目标、任务、笔记、学习路径", "支撑学生日常学习计划和进度跟踪。"),
        ("作业评测", "作业、提交、Rubric、题库、练习、错题、掌握度", "把学习结果和练习证据结构化保存。"),
        ("AI 辅助", "学生六类工作流、教师六类工作流、Provider、fallback", "把模型能力接入具体学习场景。"),
        ("协作统计", "消息、回复、提及、活动日志、SSE、统计、报告", "保留项目过程和教师分析依据。"),
    ], [1700, 3600, 4060])
    add_para(doc, "功能优先级上，身份、学习管理、作业评测和 AI 工作流属于核心功能；协作、统计、报告和运维属于增强功能。核心功能决定系统能否完成课程项目闭环，增强功能决定文档和答辩时能否展示更完整的工程边界。")

    add_para(doc, "六、UML 用例图说明", style="Heading 1")
    add_image(doc, diagrams["usecase"], "图 2 UML 用例图：学生、教师和管理员围绕不同用例访问系统。")
    add_para(doc, "学生主要使用学习、作业和 AI 辅助功能；教师主要查看学习证据、作业批改和干预建议；管理员侧重班级、成员和权限维护。当前实现中教师和管理员在部分管理入口上复用权限。")
    add_table(doc, ["参与者", "主要用例", "权限说明"], [
        ("学生", "登录系统、管理目标任务、调用 AI 学习助手、提交作业、练习复盘、查看反馈。", "只能访问自己的学习数据、AI 结果、任务草稿和提交记录。"),
        ("教师", "查看学生证据、批改作业、生成干预建议、管理题库、查看统计报告。", "可以查看教学范围内学生数据，教师 AI 草稿需要教师确认。"),
        ("管理员", "管理成员、班级、分组和角色权限，查看系统基础统计。", "当前课程项目中与教师管理入口有部分复用。"),
        ("AI 服务", "生成计划、摘要、作业指导、自检、教师草稿和报告摘要。", "作为系统服务被调用，不单独作为人类用户登录。"),
    ], [1500, 5200, 2660])
    add_para(doc, "用例图中的系统边界是 EduMind Agent。登录和权限是所有用例的前置条件；AI 相关用例依赖学习、作业或学生过程数据；教师干预和批改类用例依赖学生已产生的提交或 AI 证据。")

    add_para(doc, "七、客户端/服务端功能划分", style="Heading 1")
    add_para(doc, "客户端负责交互和展示，服务端负责权限、业务规则、数据保存和 AI 编排。这样划分可以避免把重要业务规则只放在浏览器中，也便于服务端测试。")
    add_table(doc, ["层次", "职责", "代表文件/模块"], [
        ("客户端", "登录界面、学生端/教师端路由、表单、列表、AI 面板、响应式布局、API Client 调用。", "client/src/app.js、api.js、studentRuntime.js、teacherRuntime.js、views、widgets。"),
        ("Gateway", "统一入口、静态前端托管、Token 校验、用户上下文转发、Dashboard 聚合、SSE 代理。", "services/gateway-service/src/routes.js、main.js。"),
        ("身份服务", "用户、角色、班级、小组、Token 签发和内部用户查询。", "identity-service、shared/auth/userContext.js。"),
        ("学习服务", "课程、目标、任务、笔记、学习上下文和进度计算。", "learning-service、LearningGoal、StudyTask。"),
        ("评测服务", "作业、提交、Rubric、题库、练习、错题和掌握度。", "assessment-service、assignment.js、question.js。"),
        ("AI 服务", "Prompt、Provider、学生/教师 AI 工作流、结果/草稿归档和 Provider 健康状态。", "ai-service、StudentAiWorkflowService、TeacherAiWorkspaceService。"),
        ("协作与统计服务", "消息、活动日志、SSE、统计、报告、通知、日程和运维导入。", "collaboration、analytics、report、notification、scheduler、operations。"),
    ], [1500, 4300, 3560])
    add_para(doc, "客户端不会直接读写 JSON 数据文件，也不会直接调用服务内部接口。服务端通过 Gateway 注入用户上下文，业务服务再根据角色和资源归属判断权限。")

    add_para(doc, "八、UI 截图说明", style="Heading 1")
    for index, (label, path) in enumerate(find_screenshots(), start=1):
        add_image(doc, path, f"图 {index + 2} {label}：来自项目原型截图，说明对应业务页面的布局和交互状态。")
    add_para(doc, "截图用于说明系统实际界面方向。由于旧 docs 截图已删除，本次 Word 文档优先引用 prototypes 目录下仍保留的学生端和教师端截图。")
    add_table(doc, ["截图", "说明重点", "对应需求"], [
        ("学生 AI 首页", "展示学生 AI 助手入口、学习计划和智能建议。", "R3 学生 AI 工作流。"),
        ("学习任务页面", "展示学习目标、任务进度和任务详情。", "R2 学习管理。"),
        ("作业详情页面", "展示作业要求、提交前预览和 AI 作业指导。", "R5 作业评测、R3 作业指导。"),
        ("教师 AI 工作台", "展示教师侧 AI 助手和教学建议入口。", "R4 教师 AI 工作流。"),
        ("教师批改页面", "展示作业批改、学生提交和反馈草稿。", "R4 教师批改、R5 Rubric 评分。"),
        ("协作页面", "展示课程协作消息和过程记录。", "R6 协作消息、R7 实时同步。"),
    ], [1800, 4300, 3260])
    add_para(doc, "UI 设计上，学生端更强调个人任务和 AI 行动建议；教师端更强调班级、学生、作业和右侧上下文 AI 助手；协作页面用于展示课程项目沟通和任务流转。")

    add_para(doc, "九、开发计划", style="Heading 1")
    add_para(doc, "开发计划按课程项目实际演进整理。系统不是一次性写完，而是先做最小闭环，再逐步扩展评测、AI、教师端和文档证据。")
    add_table(doc, ["阶段", "工作内容", "成员侧重点", "验收方式"], [
        ("需求与原型", "确定智能学习计划与协作系统主题，梳理目标用户、系统边界、学生端/教师端主要页面。", "王宗翰负责需求收敛；张宇航整理原型；顾昌宇评估服务可行性。", "README、prototypes、截图。"),
        ("服务骨架", "实现 Gateway、identity、learning、shared 基础设施和统一响应格式。", "顾昌宇负责后端骨架；王宗翰检查架构边界。", "服务健康检查、kernel/services 测试。"),
        ("学习主流程", "实现课程、目标、任务、笔记和进度计算。", "顾昌宇实现服务；张宇航接入页面。", "创建目标、完成任务、检查 progress。"),
        ("作业评测", "实现作业、提交、Rubric、题库、练习、错题和掌握度。", "顾昌宇负责 assessment-service；张宇航整理作业页面。", "assessment 相关服务测试。"),
        ("AI 工作流", "接入 LM Studio/Mock Provider，实现学生和教师 AI 工作区。", "顾昌宇负责 AI 服务；王宗翰审核 Prompt 边界；张宇航接入 AI 面板。", "npm run test:ai。"),
        ("协作统计", "补充协作房间、活动日志、SSE、统计、报告、通知和运维扩展。", "小组共同整理演示能力和文档依据。", "services.test.mjs。"),
        ("交付整理", "生成 UML 图、截图说明、三份 Word 文档和测试记录。", "王宗翰确认最终口径；张宇航整理文档；顾昌宇补充测试证据。", "三份 DOCX、结构校验、测试结果。"),
    ], [1200, 3600, 2600, 1960])
    add_para(doc, "计划中的每个阶段都对应可检查产物。即使 Git 提交存在集中整理情况，也可以通过开发文档、源码目录、测试文件和截图验证阶段工作是否完成。")

    add_para(doc, "十、领域类图、包图和动态 UML 说明", style="Heading 1")
    add_image(doc, diagrams["class"], "图 9 领域类图：核心对象包括 User、Course、LearningGoal、StudyTask、Assignment、Question 和 ActivityLog。")
    add_image(doc, diagrams["package"], "图 10 包图：系统按前端、网关、身份、学习、评测、AI、协作和统计报告划分逻辑包。")
    add_image(doc, diagrams["sequence"], "图 11 顺序图：AI 学习计划生成时，前端经 Gateway 调用 AI Service，并读取学习上下文。")
    add_para(doc, "动态图选择 AI 学习计划生成为例，因为它能体现前端、网关、AI 服务、学习上下文和模型 Provider 的协作关系。若模型输出异常，AI 工作流服务会使用结构化 fallback，保证页面仍能展示可用结果。")
    add_para(doc, "10.1 领域类图说明", style="Heading 2")
    add_table(doc, ["类/对象", "作用", "关系说明"], [
        ("User", "表示系统用户，包含 name、role、email、status 等基础字段。", "与课程、目标、消息、活动、AI result 都存在 owner/actor 关系。"),
        ("Course", "表示课程基础信息。", "学习目标、笔记、作业、题库和统计都通过 courseId 关联课程。"),
        ("LearningGoal", "表示学生学习目标。", "通过 recalculateProgress 根据 StudyTask 完成比例更新 progress。"),
        ("StudyTask", "表示可执行学习任务。", "属于某个 LearningGoal，可由学生确认 AI 任务草稿后创建。"),
        ("Assignment / Question", "表示作业和题库题目。", "作业提交、Rubric、练习、错题和掌握度围绕这些对象展开。"),
        ("AITutorService", "表示基础 AI 问答、计划、摘要和初评服务。", "依赖 PromptTemplate、Provider、LearningClient 和 KnowledgeClient。"),
        ("ActivityLog", "表示系统活动记录。", "用于 Dashboard、教师端过程追踪和 SSE 事件展示。"),
    ], [2100, 3800, 3460])
    add_para(doc, "10.2 包图说明", style="Heading 2")
    add_para(doc, "包图体现了前端、Gateway、身份、学习、评测、AI、协作和统计报告之间的逻辑边界。项目当前可以按多个 Node.js 服务运行，但课程文档中仍按“轻量逻辑服务划分”描述，避免把它夸大成带注册中心、消息队列和容器编排的完整微服务平台。")
    add_table(doc, ["包/服务", "输入", "输出"], [
        ("client", "用户点击、表单、路由 hash、SSE 事件。", "API 请求、页面状态、AI 面板展示。"),
        ("gateway-service", "浏览器请求和 Token。", "转发到下游服务或聚合响应。"),
        ("identity-service", "登录信息、用户和班级管理请求。", "用户上下文、Token、班级成员数据。"),
        ("learning-service", "目标、任务、笔记请求。", "学习上下文、Dashboard 学习数据、任务进度。"),
        ("assessment-service", "作业、提交、练习和评分请求。", "评测结果、错题、掌握度和学生 portfolio。"),
        ("ai-service", "学习/作业/教师上下文和 AI 请求。", "AI result、draft、actions、provider health。"),
        ("collaboration-service", "消息、内部事件和活动记录。", "协作数据、ActivityLog、SSE 推送。"),
    ], [2100, 3560, 3700])
    add_para(doc, "10.3 动态 UML 说明", style="Heading 2")
    add_para(doc, "顺序图选择 AI 学习计划生成流程，是因为它经过的模块最多：学生浏览器发起请求，Gateway 做鉴权和转发，AI Service 构造 Prompt，Learning Service 提供目标和任务上下文，最后由 LM Studio 或 Mock Provider 返回结果。这个流程能代表系统中“业务上下文 + AI 生成 + 结构化返回”的核心交互方式。")
    add_numbered(doc, [
        "学生在前端选择学习目标或打开 AI 面板。",
        "前端调用 Gateway 的 AI 接口，并携带 Bearer Token。",
        "Gateway 校验用户并转发用户上下文。",
        "AI Service 读取学习上下文并构造 Prompt。",
        "Provider 返回模型结果；如果格式不稳定，workflow 使用 fallback。",
        "前端展示计划、行动项和后续操作入口。",
    ])

    output = DOCS / "系统需求文档.docx"
    doc.save(output)
    return output


def build_design_doc(diagrams):
    doc = Document()
    title = "系统设计文档"
    design_screenshots = find_design_screenshots()
    style_document(doc, title)
    add_title_page(doc, title, "从技术选型、架构、数据、鉴权、AI 接入和部署运行角度说明系统实现。")

    add_para(doc, "一、技术选型", style="Heading 1")
    add_table(doc, ["技术项", "选择", "原因"], [
        ("前端", "浏览器原生 HTML/CSS/ES Modules", "无构建链，适合课程提交和本地演示。"),
        ("后端", "Node.js 20+", "使用内置 http、crypto、test、fs 等模块，依赖少。"),
        ("数据存储", "JSON 文件", "便于打包、查看和迁移，适合课程项目阶段。"),
        ("AI 服务", "LM Studio/OpenAI-compatible + Mock Provider", "可接入本地模型，也能在离线测试中使用 Mock。"),
        ("实时同步", "Server-Sent Events", "实现简单，适合活动日志和协作消息推送。"),
        ("测试", "Node.js Test Runner", "项目无需安装额外测试框架即可执行自动测试。"),
    ], [1700, 2800, 4860])

    add_para(doc, "二、系统架构", style="Heading 1")
    add_para(doc, "系统采用前后端分离、领域分层和逻辑服务划分的结构。当前仓库已经有 gateway、identity、learning、assessment、ai、collaboration、analytics、report、operations、knowledge、notification、scheduler 等服务目录；在课程文档中应准确表述为按逻辑服务拆分和可独立运行的轻量服务，而不是夸大为生产级云原生微服务体系。")
    add_image(doc, diagrams["package"], "图 1 包图与逻辑服务划分。")

    add_para(doc, "2.1 应用入口与整体架构对应", style="Heading 2")
    add_para(doc, "应用入口截图可以帮助说明架构图中的服务边界。前端页面由 Gateway 托管和转发，学生端、教师端看到的是不同业务工作台；后台则按身份、学习、评测、AI、协作和统计报告等逻辑服务组织。")
    add_image(doc, design_screenshots["overview"], "图 2-1 系统总览入口界面。")
    add_para(doc, "总览入口体现了系统面向学生和教师两类角色的主路径。页面本身不直接访问各个数据文件，而是通过前端 ApiClient 请求 Gateway，再由 Gateway 按路由转发到对应逻辑服务。")
    add_image(doc, design_screenshots["student_ai"], "图 2-2 学生端 AI 工作台界面。")
    add_para(doc, "学生端 AI 工作台集中展示学习计划、薄弱点、任务草稿和行动项。该界面背后主要调用 ai-service 的 student-ai 工作流，同时依赖 learning-service 提供课程、目标、任务和笔记上下文。")
    add_image(doc, design_screenshots["teacher_ai"], "图 2-3 教师端 AI 工作台界面。")
    add_para(doc, "教师端 AI 工作台面向教学计划、学生干预、作业讲评和报告摘要等场景。它与学生端共用 Provider 抽象和结构化输出机制，但权限上要求 teacher/admin 角色，并且更多依赖 assessment-service 与 analytics/report 的过程证据。")
    add_table(doc, ["界面入口", "主要后端服务", "核心数据对象", "设计说明"], [
        ("学生 AI 工作台", "gateway、ai、learning、assessment", "StudentAiResult、StudentTaskDraft、LearningGoal、StudyTask、Assignment", "AI 建议与学习任务、作业上下文绑定，结果保存后可在时间线和行动项中继续使用。"),
        ("教师 AI 工作台", "gateway、ai、assessment、analytics、report", "TeacherAiResult、TeacherAiDraft、Submission、Rubric、MasteryRecord", "AI 先生成草稿和建议，由教师确认后再进入反馈、干预或报告流程。"),
        ("协作与活动入口", "gateway、collaboration、notification", "RoomMessage、ActivityLog、CollaborationEvent、Mention", "协作消息和活动记录通过 SSE 形成轻量实时提示，页面再重新拉取最新数据。"),
        ("学习与作业入口", "gateway、learning、assessment", "Course、LearningGoal、StudyTask、Assignment、Submission", "课程学习和作业评测保持独立服务边界，页面通过统一 API 组合展示。"),
    ], [1700, 2200, 3000, 2460])

    add_para(doc, "三、子系统/逻辑服务划分", style="Heading 1")
    add_table(doc, ["逻辑服务", "主要职责"], [
        ("gateway-service", "托管前端、统一入口、鉴权转发、聚合 Dashboard 和健康检查。"),
        ("identity-service", "用户、角色、班级、小组、Token 签发和校验。"),
        ("learning-service", "课程、学习目标、任务、笔记和学习上下文。"),
        ("assessment-service", "作业、提交、Rubric、题库、练习、错题和掌握度。"),
        ("ai-service", "Prompt、Provider、学生 AI 工作流、教师 AI 工作流和 AI 调用记录。"),
        ("collaboration-service", "协作消息、活动日志、内部事件接收和 SSE 推送。"),
        ("analytics/report/operations", "统计分析、报告导出和运维导入审计等扩展能力。"),
    ], [2200, 7160])

    add_para(doc, "3.1 前端页面与后端服务对应关系", style="Heading 2")
    add_para(doc, "前端页面不是按后端服务一一拆成多个独立应用，而是在同一个浏览器端项目中按学生端和教师端组织视图。每个视图通过 ApiClient 请求 Gateway，Gateway 再把请求分发给 learning、assessment、ai 或 collaboration 等服务。")
    add_image(doc, design_screenshots["student_learn"], "图 3-1 学生学习任务页面。")
    add_para(doc, "学习任务页面主要对应 learning-service。页面展示课程、目标、任务状态和进度，服务端通过 LearningGoal、StudyTask、LearningNote 等领域对象维护学习过程；当任务完成时，还可能产生活动日志事件供 Dashboard 或协作页刷新。")
    add_image(doc, design_screenshots["teacher_assignments"], "图 3-2 教师作业管理页面。")
    add_para(doc, "教师作业管理页面主要对应 assessment-service。页面涉及 Assignment、Rubric、Submission、Question 和 PracticeSession 等对象，教师可以从这里进入发布、查看提交、批改和讲评等流程。")
    add_table(doc, ["页面或入口", "对应服务", "主要接口方向", "说明"], [
        ("学生学习任务", "learning-service", "目标、任务、笔记、学习路径", "前端把课程学习过程组织成可执行任务，后端负责进度计算和数据保存。"),
        ("学生作业与练习", "assessment-service", "作业详情、提交、自测、错题和掌握度", "作业、练习和错题保持在评测服务中，便于教师端复用同一组证据。"),
        ("学生/教师 AI 面板", "ai-service", "student-ai、teacher-ai、provider-health", "AI 结果统一归档，前端按 result、draft、action 渲染卡片和确认操作。"),
        ("教师作业管理", "assessment-service", "作业发布、批改、Rubric 和反馈", "教师侧页面不直接修改学生 AI 结果，而是读取作业和过程证据后生成反馈。"),
        ("协作空间", "collaboration-service", "消息、回复、任务、活动日志、SSE", "协作数据独立保存，避免学习和评测服务直接承担消息系统职责。"),
    ], [1800, 2100, 3000, 2460])

    add_para(doc, "四、数据存储方式", style="Heading 1")
    add_para(doc, "项目使用 data 目录下的 JSON 文件保存业务数据。每个服务拥有自己的数据文件或 collection，Repository 负责把普通记录封装成领域对象。JSON 存储便于课程阶段调试，但在真实上线时可以替换为 SQLite、MySQL 或 PostgreSQL。")
    add_table(doc, ["数据域", "典型集合"], [
        ("身份", "users、classrooms、enrollments、studyGroups、rolePermissions"),
        ("学习", "courses、goals、tasks、notes、learningPaths"),
        ("评测", "assignments、submissions、rubrics、questions、practiceSessions、mistakeItems、masteryRecords"),
        ("AI", "promptTemplates、aiRequests、aiResponses、studentAiResults、teacherAiResults"),
        ("协作", "rooms、messages、activityLogs、events"),
    ], [1600, 7760])

    add_para(doc, "五、服务协作方式", style="Heading 1")
    add_numbered(doc, [
        "浏览器只访问 Gateway 暴露的 /api/* 接口。",
        "Gateway 校验用户 Token 后，向下游服务注入 x-edumind-user-id、x-edumind-user-role、x-edumind-user-name 等上下文头。",
        "服务间内部接口使用 x-edumind-internal-key 保护，避免前端直接调用内部 API。",
        "学习、评测、AI 和协作服务通过 HTTP JSON 接口交换数据，不直接读写对方的数据文件。",
        "活动和消息通过 collaboration-service 记录，并由 SSE 推送给前端。",
    ])

    add_para(doc, "六、认证与鉴权", style="Heading 1")
    add_para(doc, "当前实现是课程项目阶段的轻量鉴权：identity-service 根据邮箱、姓名和角色进行演示登录，必要时自动创建用户，并签发 HMAC 签名的 JWT-lite Token。业务 API 通过 Authorization: Bearer token 识别用户。v12 文档中规划的 email + password、注册、sessions 和退出登录失效仍属于后续目标，本次文档不写成已完成能力。")

    add_para(doc, "七、AI 服务接入", style="Heading 1")
    add_para(doc, "AI 服务通过 Provider 抽象接入。README 中默认配置为 LM Studio 本地模型 `qwopus3.6-27b-v2-mtp@iq4_xs`，服务地址 `http://10.108.10.110:1234`，后端自动拼接 `/v1/chat/completions`。测试和离线演示可使用 Mock Provider。")
    add_table(doc, ["AI 能力", "接口或模块"], [
        ("基础问答、计划、摘要", "/api/ai/ask、/api/ai/plan、/api/ai/summarize"),
        ("学生 AI 工作流", "daily-plan、weakness-insight、task-drafts、assignment-guide、submission-check、note-organize"),
        ("教师 AI 工作流", "teaching-plan、student-intervention、assignment-commentary、feedback-draft、course-practice-plan、report-summary"),
        ("容错", "Provider 返回非 JSON 或失败时，工作流服务生成结构化 fallback。"),
    ], [2700, 6660])

    add_para(doc, "八、AI 设计", style="Heading 1")
    add_para(doc, "AI 不是单独放在页面上的聊天窗口，而是作为后端服务接入学习、作业、笔记和教师反馈流程。设计目标是：模型可以替换、输出结构稳定、结果可追踪、异常可回退，并且学生和教师看到的是贴合当前业务上下文的建议。")

    add_para(doc, "8.1 AI 总体分层", style="Heading 2")
    add_table(doc, ["层次", "职责", "对应实现"], [
        ("接口层", "暴露 /api/ai、/api/student-ai、/api/teacher-ai 等路由，负责权限和请求参数入口。", "services/ai-service/src/routes.js"),
        ("编排层", "根据 AI 类型选择学生或教师工作流，处理 result、draft、action 和 timeline。", "StudentAiWorkspaceService、TeacherAiWorkspaceService"),
        ("工作流层", "构造 Prompt，调用 Provider，解析 JSON，执行 normalize 和 fallback。", "StudentAiWorkflowService、TeacherAiWorkflowService"),
        ("领域层", "保存 PromptTemplate、AIRequest、AIResponse、ProviderHealth、StudentAiResult、TeacherAiDraft 等记录。", "services/ai-service/src/domain/ai.js"),
        ("Provider 层", "屏蔽 Mock、LM Studio、OpenAI-compatible endpoint 的差异。", "MockLLMProvider、LMStudioProvider、OpenAICompatibleProvider"),
    ], [1500, 4300, 3560])

    add_para(doc, "8.2 AI 请求流程", style="Heading 2")
    add_numbered(doc, [
        "前端在学生 AI 面板或教师 AI 面板中发起请求，例如生成每日计划、作业自检或学生干预草稿。",
        "Gateway 校验用户 Token，并把用户上下文转发给 ai-service。",
        "ai-service 根据路由选择 student-ai 或 teacher-ai 工作流，先确认用户上下文和角色权限。",
        "工作流根据当前页面传入的课程、作业、笔记、错题、学生证据等信息构造 Prompt。",
        "Provider 调用 LM Studio、OpenAI-compatible 接口或 Mock Provider。",
        "返回结果先按 JSON 解析，如果模型返回 Markdown 包裹的 JSON，也尝试提取对象。",
        "解析失败或 Provider 报错时，系统使用 fallback 模板生成结构化结果。",
        "Workspace Service 保存 AI result、draft 和 actions，供详情页、时间线和教师查看证据时复用。",
    ])

    add_para(doc, "8.3 Prompt 与上下文设计", style="Heading 2")
    add_para(doc, "Prompt 分为基础 AI 问答模板和结构化工作流 Prompt。基础问答、计划、摘要使用 PromptTemplate，根据目标、笔记和问题替换占位符；学生/教师工作流则根据当前页面上下文动态构造完整提示词。")
    add_table(doc, ["上下文来源", "使用方式"], [
        ("learning-service", "为 ask、plan、summarize 和学生工作流提供课程、目标、任务、笔记等学习上下文。"),
        ("knowledge-service", "优先提供课程知识点、学习目标、常见误区和 prompt hints；失败时回退到本地轻量资源检索。"),
        ("assessment-service", "为作业指导、自检、教师讲评、反馈草稿和补练计划提供作业、提交、Rubric、错题和掌握度信息。"),
        ("student-ai history", "教师端可以查看学生 AI results/timeline/action summary，用于生成干预和学情判断。"),
        ("前端 route state", "学生或教师当前页面会传入 route、courseId、assignmentId、studentId 等信息，AI 输出会带回对应来源。"),
    ], [2100, 7260])

    add_para(doc, "8.4 结构化输出与 Schema 归一化", style="Heading 2")
    add_para(doc, "AI 输出要求返回一个 JSON 对象，但模型实际可能返回不完整字段、纯文本或 Markdown 包裹内容。因此项目在工作流层增加 normalize 函数，把不同模型输出统一为前端能消费的结构。")
    add_table(doc, ["工作流类型", "关键结构字段"], [
        ("daily_plan", "summary、actions、risks、evidence、questions、provider、generatedAt。"),
        ("weakness_insight", "weaknesses、score、rank、evidence、action。"),
        ("task_draft", "draft.title、type、courseId、goalId、estimateMinutes、dueDate、steps、doneDefinition。"),
        ("assignment_guide", "outline、checklist、actions、evidence。"),
        ("submission_check", "completionEstimate、issues、strengths、rewriteSuggestions。"),
        ("note_organize", "cards、assignmentParagraphs，可进一步保存为学习笔记。"),
        ("teacher workflows", "summary、actions、risks、evidence、draft.title、draft.body、structuredPayload。"),
    ], [2500, 6860])

    add_para(doc, "8.5 fallback 设计", style="Heading 2")
    add_para(doc, "fallback 是本项目 AI 设计中比较重要的一点。课程演示时本地模型可能没有启动，或者模型输出不是合法 JSON。如果直接把异常暴露给页面，学生和教师端主流程会中断。因此工作流服务在 Provider 抛错、返回空内容或解析失败时，使用 fallback 结果补齐结构。")
    add_bullets(doc, [
        "Provider 抛错：记录 warning，并把 provider 标记为 fallback。",
        "返回非 JSON：保留 rawText，normalize 时使用 fallback 字段生成可展示结果。",
        "字段缺失：schema normalize 使用默认值补齐 actions、draft、issues、cards 等字段。",
        "前端收益：页面不需要判断几十种模型异常情况，只按统一结构渲染卡片、草稿和行动项。",
    ])

    add_para(doc, "8.6 学生 AI 设计", style="Heading 2")
    add_table(doc, ["学生场景", "设计意图", "落地方式"], [
        ("每日计划", "把课程目标和任务转成当天可执行安排。", "生成 daily_plan result 和 actions。"),
        ("薄弱点分析", "根据练习、错题或作业证据提示学生优先复习内容。", "生成 weakness_insight，包含 weaknesses 和证据。"),
        ("任务草稿", "把 AI 建议转成待确认学习任务。", "先生成 task_draft result，再保存 StudentTaskDraftRecord，确认后调用 learning-service 创建任务。"),
        ("作业指导", "给出作业结构、检查清单和完成建议。", "assignment_guide 输出 outline 和 checklist。"),
        ("提交前自检", "在学生提交作业前提示遗漏和修改建议。", "submission_check 输出 issues、strengths、rewriteSuggestions。"),
        ("笔记整理", "把零散笔记整理成卡片和段落。", "note_organize 可保存为正式 LearningNote。"),
    ], [1700, 3700, 3960])

    add_para(doc, "8.7 教师 AI 设计", style="Heading 2")
    add_table(doc, ["教师场景", "设计意图", "落地方式"], [
        ("教学计划", "帮助教师根据课程进度和班级情况生成下一步教学安排。", "teaching_plan 生成 result 和草稿。"),
        ("学生干预", "根据学生 AI 过程证据、作业和练习情况生成干预话术。", "student_intervention 生成 TeacherAiDraftRecord，教师确认后标记 sent。"),
        ("作业讲评", "把作业提交和 Rubric 结果整理成讲评草稿。", "assignment_commentary 输出 draft body 和 actions。"),
        ("反馈草稿", "辅助教师生成个性化批改反馈。", "feedback_draft 保存为教师草稿，确认后可保存反馈。"),
        ("课程补练计划", "根据薄弱知识点组织补练题和策略。", "course_practice_plan 输出练习策略和题目建议。"),
        ("报告摘要", "把统计和过程证据整理为报告摘要。", "report_summary 输出结构化摘要和建议动作。"),
    ], [1700, 3700, 3960])

    add_para(doc, "8.8 AI 结果存储与证据追踪", style="Heading 2")
    add_para(doc, "AI 结果不会只显示一次后丢失。学生端结果保存为 StudentAiResultRecord，教师端结果保存为 TeacherAiResultRecord，教师草稿保存为 TeacherAiDraftRecord。每条记录包含 ownerId、type、provider、result、actions、generatedAt 等字段，部分记录还关联 courseId、assignmentId、studentId、submissionId 或 sourceEvidenceIds。")
    add_bullets(doc, [
        "学生 timeline 会合并 AI result、action 状态变化和任务草稿，方便教师查看学习过程。",
        "教师 result 和 draft 分离，避免 AI 生成内容直接变成已发送反馈。",
        "action 具有 status 和 note，支持 completed、dismissed、converted、sent、saved 等状态记录。",
        "provider 和 rawText 字段便于定位当前结果来自真实模型、Mock 还是 fallback。",
    ])

    add_para(doc, "8.9 AI 权限与安全边界", style="Heading 2")
    add_table(doc, ["边界", "设计说明"], [
        ("学生只能读写自己的 AI result 和 task draft。", "StudentAiWorkspaceService 的 requireOwnedResult、requireOwnedDraft 会检查 ownerId。"),
        ("教师和管理员可以按 studentId 查看学生 AI 证据。", "resolveOwnerId 允许 teacher/admin 读取指定学生结果，普通学生跨用户访问会被拒绝。"),
        ("教师 AI 路由要求 teacher/admin 角色。", "TeacherAiWorkspaceService 通过 requireTeacher 校验角色。"),
        ("内部 AI 接口需要 internal key。", "review-submission、provider-health、student summary 等内部接口不直接暴露给普通前端调用。"),
        ("AI 输出不直接替代最终业务结果。", "作业评分、干预发送、任务创建等关键动作仍需要服务端校验或用户确认。"),
    ], [2800, 6560])

    add_para(doc, "8.10 AI 设计的局限与后续改进", style="Heading 2")
    add_bullets(doc, [
        "当前知识检索仍偏轻量，主要依赖知识服务上下文和本地资源提示，后续可以引入更完整的向量检索或知识图谱。",
        "当前 Provider 调用没有做复杂限流和队列，课程演示够用，真实上线需要增加调用额度、排队和失败重试策略。",
        "当前结构化 JSON 依赖 Prompt 约束和 normalize 兜底，后续可以增加更严格的 schema 校验和模型输出修复流程。",
        "教师端 AI 草稿已经与发送/保存动作分离，但仍需要更多人工审核提示，避免误把 AI 建议当作事实结论。",
    ])

    add_para(doc, "8.11 学生端 AI 界面设计", style="Heading 2")
    add_para(doc, "学生端 AI 界面强调“建议可以继续转化为学习动作”。系统不把模型回答当成一段普通文本展示，而是拆成摘要、证据、行动项、任务草稿和可确认操作，让学生可以从 AI 建议继续进入任务、作业或笔记流程。")
    add_image(doc, design_screenshots["student_ai_insight"], "图 8-1 学生端 AI 薄弱点分析界面。")
    add_para(doc, "薄弱点分析界面对应 student-ai 的 weakness_insight 工作流。后端根据练习、错题、作业或学习任务证据生成 weaknesses、evidence 和 actions，前端按结构化字段渲染为卡片，而不是直接显示原始模型文本。")
    add_image(doc, design_screenshots["student_note_ai"], "图 8-2 学生端 AI 笔记整理结果界面。")
    add_para(doc, "笔记整理界面对应 note_organize 工作流。AI 输出会被 normalize 为 cards 和 assignmentParagraphs，学生确认后可以保存为学习笔记；如果 Provider 不可用，fallback 仍会返回可展示的卡片结构，避免页面流程中断。")
    add_table(doc, ["学生 AI 界面元素", "后端结构", "设计意义"], [
        ("分析摘要", "result.summary、weaknesses、score、rank", "把模型判断压缩成页面可快速阅读的信息。"),
        ("证据列表", "evidence、sourceEvidenceIds、courseId、assignmentId", "说明 AI 建议来自哪些学习或评测数据，减少空泛建议。"),
        ("行动项", "actions.status、actions.note", "支持 completed、dismissed、converted 等状态，便于后续追踪。"),
        ("任务草稿", "StudentTaskDraftRecord", "AI 只生成草稿，学生确认后才调用 learning-service 创建正式任务。"),
        ("笔记卡片", "cards、assignmentParagraphs", "把模型输出转成可保存、可复习的学习资料。"),
    ], [2100, 3400, 3860])

    add_para(doc, "8.12 教师端 AI 界面设计", style="Heading 2")
    add_para(doc, "教师端 AI 界面强调“辅助教师判断和生成草稿”。教师看到的 AI 结果一般包含班级或学生证据、风险提示、建议动作和可编辑草稿，最终发送、保存或采用仍需要教师确认。")
    add_image(doc, design_screenshots["teacher_ai_plan"], "图 8-3 教师端 AI 教学计划详情界面。")
    add_para(doc, "教学计划详情界面对应 teacher-ai 的 teaching_plan 工作流。后端结合课程进度、班级掌握度和活动数据生成计划草稿，前端把 summary、actions、risks、draft 等字段拆开展示，方便教师检查后再使用。")
    add_image(doc, design_screenshots["teacher_intervention"], "图 8-4 教师端 AI 干预确认界面。")
    add_para(doc, "干预确认界面对应 student_intervention 工作流。系统会读取学生 AI timeline、作业提交、练习结果和活动记录生成干预草稿，但 TeacherAiDraftRecord 在确认前只是草稿状态，不会直接变成已发送反馈。")
    add_table(doc, ["教师 AI 流程", "输入证据", "输出结果", "人工确认点"], [
        ("教学计划", "课程目标、任务进度、班级统计", "计划摘要、风险、后续动作和草稿", "教师确认是否采用为教学安排。"),
        ("学生干预", "学生 AI 证据、作业提交、练习和活动日志", "干预话术、建议动作和原因说明", "教师确认后才标记 sent。"),
        ("作业讲评", "作业、Rubric、提交分布和典型问题", "讲评草稿和改进建议", "教师可编辑后发布或保存。"),
        ("反馈草稿", "某次提交、Rubric 评分和 AI 初评", "个性化反馈文本", "教师批改页面确认后进入反馈流程。"),
        ("报告摘要", "统计数据、过程证据和协作记录", "结构化摘要和建议", "教师决定是否写入正式报告。"),
    ], [1700, 3000, 2700, 1960])

    add_para(doc, "九、实时同步机制", style="Heading 1")
    add_para(doc, "系统使用 Server-Sent Events 实现轻量实时同步。项目没有引入 WebSocket 或消息队列，而是把协作消息、活动日志和服务内部事件集中到 collaboration-service，再通过 /api/events 向前端推送。这个方案实现成本低，适合课程项目中“有实时反馈、但不需要复杂双向长连接”的场景。")

    add_para(doc, "9.1 SSE 连接建立流程", style="Heading 2")
    add_numbered(doc, [
        "前端携带登录 Token 请求 Gateway 暴露的 /api/events。",
        "Gateway 校验 Token 后，把用户上下文和 internal key 转发给 collaboration-service。",
        "collaboration-service 检查用户上下文，通过后设置 content-type 为 text/event-stream。",
        "服务端把当前响应对象加入 SyncHub.clients 集合，并立即写入 ready 事件。",
        "浏览器保持连接，后续事件产生时由服务端主动写入 event/data 数据块。",
        "请求关闭时触发 close 回调，SyncHub 从 clients 集合中移除该连接。",
    ])

    add_para(doc, "9.2 事件产生与广播流程", style="Heading 2")
    add_table(doc, ["来源", "触发方式", "同步结果"], [
        ("学习服务", "完成任务、创建目标、更新笔记等业务动作后调用 /internal/events。", "记录 CollaborationEvent 和 ActivityLog，再向 SSE 客户端广播。"),
        ("协作服务", "发送消息、回复、创建协作任务、更新 checklist、生成 handoff 等。", "保存对应协作记录，并可形成活动日志或房间工作区更新。"),
        ("Gateway 聚合页", "Dashboard、协作页、教师工作台订阅事件。", "收到事件后可刷新活动流、协作消息或页面提示。"),
        ("测试容错场景", "下游事件服务不可用时，主业务服务捕获失败。", "主业务不回滚，事件记录失败作为可补偿问题处理。"),
    ], [1800, 4200, 3360])

    add_para(doc, "9.3 SyncHub 的实现细节", style="Heading 2")
    add_para(doc, "SyncHub 只保存当前活跃连接，不保存业务数据。connect(res) 把响应对象加入 Set，disconnect(res) 删除连接，broadcast(type, payload) 把事件写成 `event: type` 和 `data: JSON.stringify(payload)` 的标准 SSE 格式。写入某个连接失败时，服务端会把该连接移除，避免坏连接一直占用集合。")
    add_bullets(doc, [
        "clients 使用 Set 保存，避免同一个响应对象重复加入。",
        "广播时逐个写入，单个连接失败不会影响其他在线连接。",
        "SSE 是服务端到浏览器的单向推送，浏览器提交消息仍然通过普通 POST 接口完成。",
        "当前没有实现断线重放和 Last-Event-ID，刷新后由页面重新请求最新消息和活动日志恢复状态。",
    ])

    add_para(doc, "9.4 协作数据模型与同步内容", style="Heading 2")
    add_table(doc, ["对象", "用途", "同步价值"], [
        ("Room", "课程、作业、小组等协作空间。", "决定消息、任务、资源和决策归属。"),
        ("RoomMessage / MessageReply", "协作消息和回复。", "支持课程讨论、作业反馈和过程沟通。"),
        ("Mention", "消息或回复中的提醒对象。", "支持未读提醒和按用户查看提及。"),
        ("CollaborationTask / ChecklistItem", "协作任务和检查项。", "把讨论转成可跟踪事项。"),
        ("RoomSummary / RoomDecision", "阶段总结和决策记录。", "保留项目过程证据。"),
        ("SharedResource / HandoffNote", "共享资源和交接说明。", "支持课程资料流转和成员协作。"),
        ("ActivityLog / CollaborationEvent", "活动流和事件流。", "用于 Dashboard、教师端过程追踪和 SSE 推送。"),
    ], [2300, 3500, 3560])

    add_para(doc, "9.5 前端同步策略", style="Heading 2")
    add_para(doc, "前端不把 SSE 当成唯一数据来源。SSE 只负责提示“有变化”，页面仍通过普通 API 获取最新数据。这样可以避免事件丢失后状态永久不一致，也让页面刷新、重新登录和移动端恢复更简单。")
    add_table(doc, ["页面", "收到事件后的处理思路"], [
        ("总览页", "刷新 activity、任务进度、AI Provider 状态或统计摘要。"),
        ("协作页", "重新拉取房间消息、回复、提及、协作任务和 checklist。"),
        ("教师端", "刷新学生 AI 证据、干预状态、批改详情或课程活动。"),
        ("学生端", "刷新学习任务、笔记、作业反馈或 AI action 状态。"),
    ], [1800, 7560])

    add_para(doc, "9.6 实时同步的取舍", style="Heading 2")
    add_bullets(doc, [
        "优点：实现简单，浏览器原生支持，适合活动通知和协作消息推送。",
        "限制：SSE 是单向通道，不适合高频双向编辑或在线文档协同。",
        "当前取舍：课程项目优先保证可运行和可解释，因此暂不引入 WebSocket、Redis Pub/Sub 或消息队列。",
        "后续扩展：如果需要多进程部署，可把 SyncHub 后面的广播层替换为 Redis、消息队列或数据库事件表。",
    ])

    add_para(doc, "9.7 协作页面中的实时同步落点", style="Heading 2")
    add_para(doc, "协作页面是实时同步机制最直观的使用位置。消息、回复、协作任务、活动记录和提及提醒都属于 collaboration-service 的数据范围，SSE 负责通知页面“有变化”，页面再通过普通接口重新拉取最新数据。")
    add_image(doc, design_screenshots["teacher_collaboration"], "图 9-1 通用协作空间界面。")
    add_para(doc, "协作空间界面展示房间、消息和任务等内容。用户发送消息或更新协作任务后，服务端会保存 RoomMessage、CollaborationTask 或 ActivityLog，并通过 SyncHub 向在线客户端广播事件。")
    add_image(doc, design_screenshots["teacher_collaboration_detail"], "图 9-2 协作详情与过程记录界面。")
    add_para(doc, "协作详情界面体现了事件和业务数据的分离：事件只提示变化，具体消息、决策、资源、检查项和交接说明仍然保存在对应 Repository 中。这样即使 SSE 连接短暂断开，刷新页面后也能从普通 API 恢复最新状态。")
    add_table(doc, ["协作界面元素", "后端对象", "同步方式"], [
        ("房间消息", "RoomMessage、MessageReply", "消息保存后生成活动记录，并向 SSE 客户端广播 message 或 activity 类型事件。"),
        ("提及提醒", "Mention", "提及信息随消息保存，页面收到事件后刷新未读和提及列表。"),
        ("协作任务", "CollaborationTask、ChecklistItem", "任务或检查项变化后记录活动，页面刷新任务状态。"),
        ("阶段总结和决策", "RoomSummary、RoomDecision", "作为协作过程证据保存，不依赖 SSE 持久化。"),
        ("活动流", "ActivityLog、CollaborationEvent", "既用于 Dashboard 展示，也作为 SSE 广播的数据来源。"),
    ], [1800, 3000, 4560])

    add_para(doc, "十、日志与故障定位", style="Heading 1")
    add_para(doc, "系统的故障定位主要依靠四类信息：统一错误响应、服务健康检查、活动/审计记录和测试中覆盖的异常场景。项目没有引入 ELK、Prometheus 等生产级监控组件，但在课程项目范围内保留了可以定位问题的结构化线索。")

    add_para(doc, "10.1 统一错误响应", style="Heading 2")
    add_para(doc, "shared/http/errors.js 定义了 AppError、ValidationError、AuthError、ForbiddenError 和 NotFoundError。所有异常最终转换为统一 JSON 结构，包含 ok=false、code、message 和 details。前端和测试都可以根据 code 判断错误类型。")
    add_table(doc, ["错误码", "HTTP 状态", "典型场景"], [
        ("VALIDATION_ERROR", "400", "请求体不是合法 JSON、必填字段为空、状态流转参数不合法。"),
        ("AUTH_REQUIRED", "401", "缺少 Token、Token 签名错误或登录状态过期。"),
        ("FORBIDDEN", "403", "学生访问教师专属 AI 路由、跨用户读取学生 AI 证据。"),
        ("NOT_FOUND", "404", "访问不存在的任务、笔记、AI result、草稿或协作对象。"),
        ("DOWNSTREAM_UNAVAILABLE", "503", "Gateway 或服务 client 无法访问下游服务。"),
        ("DOWNSTREAM_TIMEOUT", "504", "内部 ServiceClient 请求超过 timeoutMs。"),
        ("INTERNAL_ERROR", "500", "未预期异常，返回通用服务不可用信息。"),
    ], [2500, 1500, 5360])

    add_para(doc, "10.2 ServiceClient 的下游故障处理", style="Heading 2")
    add_para(doc, "Gateway 和各服务之间通过 shared/client/ServiceClient 调用 HTTP JSON 接口。ServiceClient 会设置超时、解析下游 JSON、识别 ok=false 响应，并把连接失败、超时、非法响应统一转换成 AppError。这样上层路由不用分别处理 fetch failed、AbortError、JSON parse error 等底层异常。")
    add_bullets(doc, [
        "下游服务未配置 baseUrl 时，返回 DOWNSTREAM_UNAVAILABLE。",
        "请求超时时，AbortController 触发 DOWNSTREAM_TIMEOUT。",
        "下游返回非 JSON 时，标记 INVALID_RESPONSE。",
        "下游业务错误 ok=false 时，保留下游 code 和 message。",
        "普通网络失败时，转换为 DOWNSTREAM_UNAVAILABLE，并在 details 中记录 service 和 cause。",
    ])

    add_para(doc, "10.3 活动日志与事件记录", style="Heading 2")
    add_para(doc, "ActivityService 同时维护 ActivityLog 和 CollaborationEvent。recordEvent 先保存事件，再创建活动日志，最后调用 SyncHub.broadcast。活动日志用于给用户看“发生了什么”，事件记录用于服务间同步和后续分析。")
    add_table(doc, ["记录类型", "主要字段", "用途"], [
        ("ActivityLog", "actorId、type、summary、payload、createdAt", "展示最近活动、Dashboard 活动流、教师查看学生过程。"),
        ("CollaborationEvent", "type、source、actorId、summary、payload、occurredAt", "记录从其他服务传来的领域事件，并触发 SSE。"),
        ("AuditRecord", "actorId、action、resourceType、before、after、metadata、occurredAt", "协作模块中更细的操作审计，例如资源、任务、交接变更。"),
        ("ProviderHealthRecord", "provider、model、status、endpoint、checkedAt", "定位 AI 当前使用 Mock、LM Studio 还是 OpenAI-compatible Provider。"),
    ], [2100, 3860, 3400])

    add_para(doc, "10.4 健康检查与定位入口", style="Heading 2")
    add_para(doc, "每个服务都有 /health 接口，Gateway 额外提供 /api/health 聚合下游状态。聚合结果包含服务名称、status、latencyMs 和错误信息。当前实现中，只要有一个下游服务不是 up，Gateway 整体状态就会变成 degraded。")
    add_numbered(doc, [
        "先访问 Gateway 的 /api/health，判断是整体可用、部分 degraded 还是入口服务异常。",
        "如果某个下游服务显示 down，再单独访问该服务 /health，确认是服务未启动、端口错误还是请求超时。",
        "如果 AI 功能异常，查看 /internal/ai/provider-health 或相关测试，确认 Provider、model 和 endpoint。",
        "如果协作或实时同步异常，检查 collaboration-service /health、/api/activity 和 /api/events。",
    ])

    add_para(doc, "10.5 常见故障定位表", style="Heading 2")
    add_table(doc, ["现象", "优先检查", "可能原因"], [
        ("登录后 /api/me 失败", "identity-service /health、Token 是否携带、internal key 是否一致。", "Token 缺失、签名错误、identity-service 未启动。"),
        ("Dashboard 加载失败", "Gateway 聚合日志、learning/assessment/collaboration/ai/analytics 健康状态。", "某个下游服务不可用或返回 ok=false。"),
        ("AI 页面没有结果", "ai-service /health、ProviderHealth、LM Studio 地址、Mock Provider 配置。", "模型服务未启动、endpoint 拼接错误、Provider 超时。"),
        ("AI 卡片字段不完整", "ai-schemas normalize 和 rawText。", "模型没有按 JSON 输出，系统使用 fallback 补齐。"),
        ("协作消息不刷新", "collaboration-service /api/events、SyncHub 连接、浏览器是否断线。", "SSE 连接断开、用户上下文失效、页面未重新拉取数据。"),
        ("任务完成但活动流没更新", "learning-service 事件发布、/internal/events、activityLogs。", "事件服务短暂不可用；主业务完成但事件记录失败。"),
        ("教师无法查看学生 AI 证据", "用户角色、studentId、teacher-ai route guard。", "普通学生跨用户访问，或教师上下文未正确转发。"),
    ], [2200, 3860, 3300])

    add_para(doc, "10.6 测试对故障定位的支撑", style="Heading 2")
    add_para(doc, "测试代码不是只验证成功路径，也专门覆盖故障路径。services.test.mjs 中包含下游不可用、Gateway degraded、协作事件失败不回滚等场景；ai-route-guards.test.mjs 覆盖缺少用户上下文、内部 key 错误和角色越权；ai-workflows.test.mjs 覆盖 Provider 返回非法 JSON 或抛错时的 fallback。")
    add_bullets(doc, [
        "当接口问题复现不稳定时，优先找对应测试文件，看预期错误码和边界条件。",
        "当模型输出异常时，优先看 ai-workflows 和 ai-schemas 测试，确认是否应走 fallback。",
        "当服务调用失败时，优先看 services.test 中 Gateway 和 ServiceClient 的异常处理测试。",
        "当权限问题出现时，优先看 ai-route-guards、identity-service 和 userContext 相关测试。",
    ])

    add_para(doc, "10.7 当前日志体系的不足", style="Heading 2")
    add_bullets(doc, [
        "当前日志主要保存在 JSON 数据和测试输出中，没有集中日志平台。",
        "服务端没有实现请求 ID 或 trace ID，因此跨服务追踪还需要根据时间、用户和业务 ID 人工关联。",
        "SSE 连接数、Provider 调用耗时和失败率没有形成长期指标。",
        "后续如果真实部署，可以补充结构化 console 日志、requestId、日志文件轮转和简单监控面板。",
    ])

    add_para(doc, "十一、部署方式、运行环境、安装与卸载", style="Heading 1")
    add_table(doc, ["项目", "说明"], [
        ("运行环境", "Windows 或其他支持 Node.js 20+ 的环境；浏览器建议 Chrome/Edge。"),
        ("启动", "单体入口可执行 npm start；逻辑服务模式可执行 npm run start:services:mock 或 npm run start:services。"),
        ("访问", "默认 http://127.0.0.1:4077。"),
        ("安装", "解压项目并准备 Node.js；项目原则上不需要 npm install。"),
        ("卸载", "删除项目目录即可；运行数据主要位于 data 目录。"),
        ("AI 模型", "如需真实 AI，需在 LM Studio 启动 OpenAI-compatible Server。"),
    ], [1800, 7560])

    add_para(doc, "十二、技术关键点展开", style="Heading 1")
    add_para(doc, "本节说明项目中比较关键的技术实现。描述重点放在源码中已经出现的设计和流程，不把后续规划写成已经上线的能力。")

    add_para(doc, "12.1 请求进入系统后的处理链路", style="Heading 2")
    add_para(doc, "浏览器端通过 ApiClient 访问 Gateway 暴露的 /api/* 接口。Gateway 先解析 Authorization Bearer Token，调用 identity-service 校验用户，再把用户 ID、角色和名称写入内部请求头。下游服务只读取这些受信任的用户上下文头，不要求每个服务重复解析浏览器 Token。")
    add_table(doc, ["步骤", "处理位置", "实现要点"], [
        ("1", "client/src/api.js", "统一封装 request、query 参数、JSON body 和错误响应处理。"),
        ("2", "gateway-service", "校验登录态，转发请求，聚合 Dashboard、报告、AI、统计等下游结果。"),
        ("3", "shared/auth/userContext.js", "定义用户上下文头和内部 key 校验逻辑。"),
        ("4", "业务服务 routes.js", "只做参数读取、用户上下文获取和应用服务调用，避免把业务规则写在路由层。"),
        ("5", "application/domain/infrastructure", "应用服务编排用例，领域对象处理核心规则，Repository 负责 JSON 数据读写。"),
    ], [700, 2700, 5960])

    add_para(doc, "12.2 服务内部的分层结构", style="Heading 2")
    add_para(doc, "每个主要业务服务都尽量保持 main、config、routes、application、domain、infrastructure 的结构。这样做的好处是：入口、HTTP、业务流程、领域对象和外部依赖分开，后续如果要替换存储或拆成更独立的服务，修改范围比较清楚。")
    add_table(doc, ["层次", "职责", "例子"], [
        ("main/config", "启动 HTTP 服务、读取端口、内部 key、Provider 等配置。", "services/ai-service/src/main.js、config.js"),
        ("routes", "定义 HTTP 路径，读取 req body/query/params，返回统一 ok/error 结构。", "services/assessment-service/src/routes.js"),
        ("application", "编排一个完整用例，处理权限、状态流转、跨服务调用和事件发布。", "learningService、gradingService、studentAiWorkspaceService"),
        ("domain", "表达业务对象和核心方法，不直接依赖 HTTP 和文件系统。", "LearningGoal、StudyTask、Assignment、Question、MasteryRecord"),
        ("infrastructure", "封装 JSON 数据、种子数据、服务 client、Provider 和知识检索。", "seed.js、clients/*、resourceSearch.js"),
    ], [1600, 4300, 3460])

    add_para(doc, "12.3 领域对象与 Repository 的使用", style="Heading 2")
    add_para(doc, "项目没有把所有数据都当成普通 JSON 操作，而是通过 Entity 和 Repository 包装业务对象。Repository 负责从指定 collection 读取记录并转换成领域对象，应用服务再调用领域对象方法完成业务规则。")
    add_bullets(doc, [
        "LearningGoal.recalculateProgress(tasks) 根据同一目标下任务完成比例计算 progress，并在 100% 时把目标状态改为 completed。",
        "StudyTask.complete() 只负责把任务状态改为 done 并更新时间，具体谁能完成任务由应用服务判断。",
        "PracticeSession、AnswerRecord、MistakeItem、MasteryRecord 分别对应练习、答题、错题和知识点掌握记录，便于后续统计分析复用。",
        "跨服务引用只保存 ID 或展示快照，例如 submission 保存 studentId 和 studentSnapshot，权限判断仍以身份服务查询结果为准。",
    ])

    add_para(doc, "12.4 数据一致性与失败处理", style="Heading 2")
    add_para(doc, "当前项目使用 JSON 文件，不引入分布式事务。写操作由所属服务先完成本地数据修改，再向协作服务发布活动事件。事件失败时不回滚主业务，而是在日志和测试中明确作为可容错场景处理。")
    add_table(doc, ["场景", "处理方式", "原因"], [
        ("完成学习任务", "learning-service 更新 task，重新计算 goal progress，再尝试发布 activity/event。", "任务完成是主业务，事件用于通知和展示。"),
        ("AI 初评失败", "assessment-service 调用 ai-service 失败时返回明确错误或保留待处理状态。", "避免把不可靠 AI 结果写成最终评分。"),
        ("Provider 输出非法 JSON", "AI workflow 尝试解析失败后返回结构化 fallback。", "保证前端卡片和草稿页面仍能展示。"),
        ("下游服务不可用", "Gateway 或 ServiceClient 返回 DOWNSTREAM_UNAVAILABLE 等统一错误。", "前端可以获得可解释错误，而不是空白页面。"),
    ], [1900, 4100, 3360])

    add_para(doc, "12.5 AI Provider 与结构化工作流", style="Heading 2")
    add_para(doc, "AI 服务没有直接把模型回答透传给前端，而是拆成 Prompt 构造、Provider 调用、结构化解析、fallback 和结果归档几个步骤。这样可以兼容 LM Studio、OpenAI-compatible endpoint 和 Mock Provider，也方便测试模型异常输出。")
    add_numbered(doc, [
        "Prompt 构造阶段：根据学生或教师当前页面上下文，拼接课程、任务、作业、笔记、错题和历史 AI 结果。",
        "Provider 调用阶段：根据环境配置选择 Mock、LM Studio 或 OpenAI-compatible Provider。",
        "解析阶段：优先解析模型返回的 JSON；如果模型返回 Markdown 包裹 JSON，也尝试提取其中结构。",
        "归一化阶段：把不完整字段补齐为前端可用的卡片、行动项、草稿或报告摘要。",
        "落库阶段：studentAiWorkspace 或 teacherAiWorkspace 保存 result、draft、action 状态，供列表、详情和时间线复用。",
    ])

    add_para(doc, "12.6 学生端与教师端 AI 流程差异", style="Heading 2")
    add_table(doc, ["端", "主要入口", "设计重点"], [
        ("学生端", "daily-plan、weakness-insight、task-drafts、assignment-guide、submission-check、note-organize", "帮助学生把学习活动转成可执行计划、草稿、自检和笔记，不直接代替学生提交。"),
        ("教师端", "teaching-plan、student-intervention、assignment-commentary、feedback-draft、course-practice-plan、report-summary", "帮助教师理解学生过程证据，生成干预和反馈草稿，由教师确认后再使用。"),
        ("共同点", "result、draft、action、timeline", "结果可追踪、动作可更新、草稿可确认，便于前端稳定展示。"),
    ], [1400, 4300, 3660])

    add_para(doc, "12.7 SSE 实时同步实现", style="Heading 2")
    add_para(doc, "协作服务中的 SyncHub 维护已连接的响应对象集合。浏览器访问 /api/events 后保持 text/event-stream 连接，服务端在消息、活动或内部事件产生时调用 broadcast，把事件类型和 payload 写入每个连接。该方式比轮询更实时，比 WebSocket 更简单，适合本项目的课程协作和活动提醒。")
    add_bullets(doc, [
        "连接建立：服务端把 response 加入 clients 集合，并设置 SSE 所需响应头。",
        "事件广播：broadcast(type, payload) 生成 event/data 格式文本并写入所有连接。",
        "异常清理：如果某个连接写入失败，服务端从集合中移除该连接。",
        "适用范围：活动日志、协作消息、任务变化、Dashboard 刷新提示等轻量实时通知。",
    ])

    add_para(doc, "12.8 前端路由、状态和视图组织", style="Heading 2")
    add_para(doc, "前端使用原生 ES Modules，而不是 React/Vue。app.js 负责登录、全局状态、hash 路由和事件分发；studentRuntime、teacherRuntime 分别处理学生端和教师端默认路由与 shell；views、forms、widgets、state、ai、utils 目录负责拆分页面、表单、组件、选择器和 AI 适配器。")
    add_table(doc, ["目录/文件", "设计作用"], [
        ("client/src/api.js", "统一封装前端接口调用，减少页面直接拼接口的重复代码。"),
        ("client/src/state", "把原始后端数据转换为页面需要的 view model。"),
        ("client/src/views/student", "学生端学习、作业、练习、笔记、AI、反馈等页面。"),
        ("client/src/views/teacher", "教师端首页、课程、学生画像、作业批改、干预和报告页面。"),
        ("client/src/ai", "学生/教师 AI Adapter 和前端 schema，处理正式 API 与 fallback。"),
        ("client/src/widgets/forms", "复用布局、表格、卡片、AI 面板和表单输入。"),
    ], [2400, 6960])

    add_para(doc, "12.9 统一错误响应与测试验证", style="Heading 2")
    add_para(doc, "shared/http 中定义了统一响应和错误结构。正常响应返回 ok/data/meta，异常响应返回 ok=false、code、message 和 details。测试中覆盖了登录、服务健康、内部 key、跨用户访问、Provider fallback、下游服务不可用和 Gateway 转发等场景，用来验证接口边界是否稳定。")

    add_para(doc, "12.10 作业提交与批改流程界面", style="Heading 2")
    add_para(doc, "作业流程贯穿学生端和教师端。学生端负责查看作业要求、提交内容和接收反馈；教师端负责查看提交、Rubric 评分、AI 初评建议和最终反馈。两端共享 assessment-service 中的 Assignment、Submission、Rubric 和 Feedback 等数据对象。")
    add_image(doc, design_screenshots["student_assignment_preview"], "图 12-1 学生端作业提交预览界面。")
    add_para(doc, "学生端作业预览界面对应作业详情和提交前检查流程。页面读取 Assignment、Rubric 和已有 Submission，并可调用 student-ai 的 assignment_guide 或 submission_check 工作流生成检查清单，但最终提交仍由 assessment-service 保存。")
    add_image(doc, design_screenshots["teacher_grading"], "图 12-2 教师端作业批改详情界面。")
    add_para(doc, "教师端批改详情界面对应 assessment-service 的评分和反馈流程。AI 初评可以作为参考信息进入页面，但教师评分、Rubric 结论和反馈发布仍由教师确认，避免把模型输出直接当作最终成绩。")
    add_table(doc, ["流程步骤", "涉及服务", "关键对象", "设计要点"], [
        ("作业查看", "assessment-service", "Assignment、Rubric", "学生端读取作业要求、截止时间、评分标准和提交状态。"),
        ("提交前自检", "ai-service、assessment-service", "SubmissionCheckResult、Submission", "AI 根据作业要求和草稿内容给出遗漏提示，结果只作为辅助建议。"),
        ("正式提交", "assessment-service", "Submission", "服务端保存提交记录和学生快照，后续教师批改复用同一份证据。"),
        ("AI 初评", "ai-service、assessment-service", "AIResponseRecord、Rubric", "AI 结合 Rubric 生成初步反馈或问题提示，异常时返回明确错误或 fallback。"),
        ("教师确认", "assessment-service", "Grade、Feedback", "教师确认评分和反馈后才形成最终批改结果。"),
        ("活动记录", "collaboration-service", "ActivityLog、CollaborationEvent", "提交、批改或反馈发布可形成活动流，供 Dashboard 和协作页展示。"),
    ], [1400, 2300, 2500, 3160])

    add_para(doc, "12.11 数据分析与报告界面", style="Heading 2")
    add_para(doc, "数据分析和报告功能用于把学习任务、作业结果、练习掌握度、AI 过程证据和协作活动整理成教师可读的信息。当前实现属于课程项目阶段的统计分析和报告导出能力，不写成完整的数据仓库或商业 BI 平台。")
    add_image(doc, design_screenshots["teacher_analytics"], "图 12-3 报告与导出界面。")
    add_para(doc, "报告与导出界面主要对应 report-service，并使用 analytics-service、learning-service、assessment-service、ai-service 和 collaboration-service 形成的过程数据。页面把课程周报、作业评阅、AI 使用报告等材料组织成可生成、可预览和可导出的报告条目。")
    add_table(doc, ["分析内容", "数据来源", "实现意义"], [
        ("学习进度", "LearningGoal、StudyTask、LearningPath", "反映课程目标完成情况和学生当前任务推进状态。"),
        ("作业表现", "Assignment、Submission、Rubric、Grade", "帮助教师识别作业中的共性问题和个体差异。"),
        ("练习掌握度", "PracticeSession、AnswerRecord、MistakeItem、MasteryRecord", "把错题和知识点掌握情况转成后续补练依据。"),
        ("AI 过程证据", "StudentAiResult、TeacherAiResult、AIRequestRecord", "辅助判断学生是否使用过计划、自检、笔记整理等 AI 支持。"),
        ("协作活动", "ActivityLog、RoomMessage、CollaborationTask", "展示小组协作和课程过程记录，支持项目文档中的过程证明。"),
        ("报告输出", "Report、ReportSection、summary payload", "把统计结论整理为可提交或可查看的报告摘要。"),
    ], [2100, 3500, 3760])

    output = DOCS / "系统设计文档.docx"
    doc.save(output)
    return output


def build_management_doc(output_path: Path | None = None):
    doc = Document()
    title = "项目管理文档"
    style_document(doc, title)
    add_title_page(doc, title, "记录成员分工、AI 使用、开发过程和测试记录。")

    add_para(doc, "一、成员分工", style="Heading 1")
    add_para(doc, "本项目为三人小组完成。以下分工按课程项目交付口径整理，用于说明每名成员主要负责的方向；具体贡献占比建议由小组在最终提交前共同确认，不在文档中编造精确百分比。")
    add_table(doc, ["成员", "角色", "主要工作"], [
        ("王宗翰", "组长", "统筹项目主题、需求拆解、最终集成、文档口径确认和提交材料检查。"),
        ("顾昌宇", "组员", "参与后端服务、领域模型、AI 接入、数据结构和接口测试整理。"),
        ("张宇航", "组员", "参与前端页面、原型截图、测试记录、UML/文档材料整理。"),
    ], [1300, 1500, 6560])

    add_para(doc, "二、贡献占比说明", style="Heading 1")
    add_para(doc, "本项目以小组协作方式完成，开发过程中存在共同讨论、交叉修改和统一整理材料的情况。贡献占比不按单一文件数量或提交次数计算，而是按需求分析、架构实现、前端交互、AI 接入、测试和文档整理等实际工作综合确认。")
    add_table(doc, ["成员", "贡献方向", "贡献占比"], [
        ("王宗翰", "项目统筹、需求拆解、架构取舍、最终集成、文档口径和提交材料检查。", "待小组确认"),
        ("顾昌宇", "后端服务、领域模型、数据结构、AI 接入、服务接口和测试用例整理。", "待小组确认"),
        ("张宇航", "前端页面、原型截图、交互流程、测试记录、UML 图和课程文档整理。", "待小组确认"),
    ], [1300, 5660, 2400])

    add_para(doc, "三、AI 应用情况", style="Heading 1")
    add_para(doc, "本项目既使用 AI 辅助开发，也把 AI 能力做进系统业务流程中。小组在使用 AI 时没有把生成内容直接作为最终结果，而是结合源码、测试和运行效果进行筛选、修改和复核。")
    add_table(doc, ["应用环节", "具体做法", "人工复核方式"], [
        ("需求分析", "将课程作业要求、README 和已有开发文档输入 AI，辅助整理系统边界、目标用户、功能分解和三份课程文档的大纲。", "组长对照课程要求逐项检查，避免漏掉可行性、UML、部署、测试和项目管理等章节。"),
        ("架构讨论", "让 AI 对前后端分离、分层单体、逻辑服务划分和后续微服务边界进行比较，辅助确定轻量 Node.js 服务结构。", "结合 docs/dev/architecture.md 和 service-boundaries.md 修正文案，只写已实现或有明确边界的内容。"),
        ("后端开发", "使用 AI 辅助生成身份、学习、评测、AI、协作等服务的路由草稿、领域对象和错误处理思路。", "顾昌宇结合实际源码逐段调整，重点检查接口路径、权限判断、Repository 数据归属和 JSON 响应格式。"),
        ("前端开发", "使用 AI 辅助拆分学生端、教师端页面结构，生成部分视图、表单和状态选择器代码。", "张宇航通过浏览器截图、前端测试和页面交互检查布局是否可用，删除不符合项目风格的描述性文案。"),
        ("AI Prompt 设计", "使用 AI 辅助设计学生每日计划、薄弱点分析、任务草稿、作业指导、自检、笔记整理和教师干预等 Prompt 输出结构。", "小组把输出限制为结构化 schema，并通过 ai-schemas、ai-workflows、ai-workspaces 测试验证 fallback。"),
        ("测试补全", "使用 AI 根据服务边界和异常场景列出测试点，如下游不可用、Provider 输出非法 JSON、跨用户访问、内部 key 错误等。", "最终测试写入 test/*.mjs，执行 npm test、npm run test:services 和 npm run test:ai 确认通过。"),
        ("文档写作", "使用 AI 辅助把技术实现转写为中文课程项目文档，生成需求、设计和管理文档中的说明段落。", "只保留能从 README、源码、测试、截图或开发文档找到依据的内容；对规划项明确写成后续目标。"),
    ], [1600, 4300, 3460])
    add_para(doc, "系统内置 AI 功能", style="Heading 2")
    add_table(doc, ["用户", "AI 功能", "对应实现"], [
        ("学生", "每日学习计划、薄弱点分析、任务草稿、作业指导、提交前自检、笔记整理。", "ai-service 的 student-ai 工作流、StudentAiWorkflowService、StudentAiWorkspaceService。"),
        ("教师", "教学计划、学生干预建议、作业讲评、反馈草稿、课程补练计划、报告摘要。", "ai-service 的 teacher-ai 工作流、TeacherAiWorkflowService、TeacherAiWorkspaceService。"),
        ("系统", "Provider 健康检查、Mock fallback、LM Studio/OpenAI-compatible endpoint 归一化、AI 请求和响应记录。", "AITutorService、Provider 配置、PromptTemplate、AI 仓储和相关测试。"),
    ], [1200, 4300, 3860])
    add_para(doc, "AI 使用边界", style="Heading 2")
    add_bullets(doc, [
        "AI 生成内容主要作为草稿、建议和辅助判断，不直接替代教师最终评分或学生最终提交。",
        "当模型不可用或返回格式不稳定时，系统使用 Mock Provider 或结构化 fallback 保证页面可演示。",
        "涉及账号、权限、服务边界、测试结果和代码量等事实信息时，以源码和命令输出为准，不以 AI 回答为准。",
    ])

    add_para(doc, "四、实际开发记录", style="Heading 1")
    add_para(doc, "实际开发记录按开发文档中的版本阶段和小组分工整理，不按提交记录逐条对应。每个阶段都包含任务目标、成员参与和可复查产物。")
    add_table(doc, ["阶段", "主要工作", "成员参与", "产出或依据"], [
        ("需求与边界确认", "讨论系统主题，确定面向课程学习计划、AI 辅导和协作记录，不覆盖正式教务、支付和生产级账号体系。", "王宗翰负责需求收敛；顾昌宇补充后端可行性；张宇航整理原型方向。", "README、docs/dev/versions/v0-current-monolith、系统需求文档。"),
        ("服务骨架与架构规划", "确定前后端分离、领域分层和逻辑服务划分，明确 gateway、identity、learning、assessment、ai、collaboration 等服务边界。", "王宗翰负责架构取舍；顾昌宇搭建服务目录和 shared 基础设施；张宇航确认前端调用方式。", "docs/dev/architecture.md、docs/dev/service-boundaries.md、services/、shared/。"),
        ("身份与学习主流程", "实现演示登录、Token 校验、用户角色、课程、学习目标、学习任务、学习笔记和进度计算。", "顾昌宇实现身份和学习服务；王宗翰检查业务边界；张宇航接入前端学习页面。", "identity-service、learning-service、LearningGoal.recalculateProgress、kernel/client/services 测试。"),
        ("作业评测与练习模块", "补充作业、提交、Rubric、题库、练习会话、答题记录、错题本和掌握度记录。", "顾昌宇负责 assessment-service 领域模型和应用服务；张宇航整理作业、练习、错题页面；王宗翰做功能验收。", "assessment-service、question.js、assignment.js、practiceService、services.test.mjs。"),
        ("AI 服务与 Prompt 工作流", "接入 Mock Provider 和 LM Studio/OpenAI-compatible Provider，设计学生端六类 AI 工作流和教师端六类 AI 工作流。", "顾昌宇负责 ai-service 和 Provider；王宗翰审核 Prompt 输出边界；张宇航把结果接入学生/教师 AI 面板。", "ai-service、studentAiWorkflowService、teacherAiWorkflowService、ai-schemas/workflows/workspaces 测试。"),
        ("学生端 AI-first 前端", "根据 student-ai-first-v2 原型完善学生端学习、作业、练习、笔记、AI 洞察和提交自检流程。", "张宇航负责页面和交互整理；顾昌宇补齐对应 API；王宗翰检查是否符合课程项目主线。", "client/src/views/student、studentRuntime、prototypes/student-ai-first-v2、client.test.mjs。"),
        ("教师端 AI-first 前端", "根据 teacher-ai-first 原型完善教师首页、课程、学生画像、批改、干预、报告和右侧上下文 AI 助手。", "张宇航负责教师端 shell 和页面层级；顾昌宇提供学生 AI 证据接口；王宗翰做角色权限和信息展示检查。", "client/src/views/teacher、teacherRuntime、prototypes/teacher-ai-first、v10/v11 开发文档。"),
        ("协作、通知、统计和报告扩展", "补充协作消息、活动日志、SSE、通知提醒、统计分析、报告导出和运维导入审计能力。", "顾昌宇实现服务端能力；张宇航接入工作台入口；王宗翰检查文档中是否准确描述为课程项目扩展能力。", "collaboration-service、notification-service、scheduler-service、analytics-service、report-service、operations-service。"),
        ("测试与交付整理", "补充自动测试、代码量统计、截图/UML 说明和三份课程 Word 文档。", "王宗翰负责最终口径和提交检查；顾昌宇补齐服务/AI 测试说明；张宇航整理 UI 截图、测试记录和文档排版。", "npm test、npm run test:services、npm run test:ai、linecount、三份 Word 文档。"),
    ], [1300, 3350, 2450, 2260])
    add_para(doc, "阶段复盘说明", style="Heading 2")
    add_bullets(doc, [
        "开发过程中先保证学习、作业、AI 和协作主流程能跑通，再补充统计、报告、通知和运维等扩展能力。",
        "学生端和教师端前端不是一次性完成，而是先做原型，再把稳定的信息架构迁移到真实 API 和状态管理中。",
        "AI 工作流是项目后期重点，原因是它同时影响前端展示、后端 Prompt、Provider 容错和测试用例。",
        "最终文档以开发文档、源码目录和测试结果为依据，避免把规划中的真实密码账号和生产级部署写成已完成。",
    ])

    add_para(doc, "五、测试记录", style="Heading 1")
    add_para(doc, "测试记录按照实际测试代码和本次命令执行结果整理。项目使用 Node.js 内置 test runner，测试文件集中在 test 目录；不同命令覆盖范围不同，不能只用一行“测试通过”概括。")
    add_table(doc, ["测试命令", "本次结果", "覆盖范围"], [
        ("npm test", "82/82 通过", "课程提交默认回归测试，覆盖核心内核、前端视图模型、AI schema、AI workflow、AI workspace、Prompt、客户端集成和学生/教师路由。"),
        ("npm run test:services", "32/32 通过", "服务级集成测试，覆盖各逻辑服务健康检查、Gateway 转发、身份、学习、评测、AI、协作、通知、统计、报告和运维接口。"),
        ("npm run test:ai", "65/65 通过", "AI 专项测试，额外覆盖 ai-service 路由、AI route guard、学生/教师 AI 工作区生命周期和权限边界。"),
        ("npm run linecount", "48335 total", "代码量统计，排除文档、截图、日志、压缩包、生成文件和运行数据。"),
    ], [2100, 1800, 5460])

    add_para(doc, "5.1 测试文件覆盖明细", style="Heading 2")
    add_table(doc, ["测试文件", "用例数", "主要验证内容"], [
        ("kernel.test.mjs", "6", "登录签发 Token、学习目标/任务进度计算、Mock AI、LM Studio Provider 配置、OpenAI-compatible reasoning_content 兼容。"),
        ("client.test.mjs", "21", "前端格式化、表单校验、ApiClient 路径、学生/教师路由、教师 v11 shell、AI 面板、学生详情 XSS 转义。"),
        ("ai-schemas.test.mjs", "15", "学生每日计划、薄弱点、任务草稿、作业指导、自检、笔记整理，以及教师计划类结果的结构化归一化和 fallback。"),
        ("ai-workflows.test.mjs", "10", "学生/教师 AI Workflow 的用户上下文校验、JSON/Markdown JSON 解析、Provider 异常 fallback、六类工作流分发。"),
        ("ai-workspaces.test.mjs", "10", "学生 AI result 保存、列表、详情、action 更新、时间线、任务草稿、笔记保存，以及教师 result/draft 生命周期和权限。"),
        ("ai-client-integration.test.mjs", "6", "StudentAiAdapter、TeacherAiAdapter 优先调用正式 API，并在接口不可用时回退到 askAI/summarizeNote。"),
        ("ai-prompts-domain.test.mjs", "6", "Prompt 构造、PromptTemplate 占位符替换、LM Studio endpoint 归一化、学生/教师 AI 仓储过滤与排序。"),
        ("ai-selector-views.test.mjs", "8", "DOM 转义、学生/教师共享视图稳定渲染、不同 route 下 selector 和 AI panel 上下文。"),
        ("ai-service-routes.test.mjs", "4", "ai-service 学生/教师 AI 路由、note organize、教师查看学生证据、内部 provider health 和 submission review。"),
        ("ai-route-guards.test.mjs", "6", "公共 AI 路由缺少用户上下文、内部 key 错误、教师专属路由、跨用户读取/编辑、未知草稿和权限错误。"),
        ("services.test.mjs", "32", "服务健康、Gateway 聚合、ServiceClient、知识库、评测、身份、学习、AI、协作、通知、统计、报告、运维等服务级流程。"),
    ], [2600, 900, 5860])

    add_para(doc, "5.2 服务级测试用例明细", style="Heading 2")
    add_table(doc, ["序号", "测试点", "验收含义"], [
        ("1", "各服务 health endpoint 返回 ok", "确认 gateway、identity、learning、assessment、ai、collaboration 等服务能独立启动并暴露健康状态。"),
        ("2", "Gateway 聚合下游服务状态", "确认统一入口能汇总服务健康结果。"),
        ("3", "下游服务不可用时 Gateway 仍返回 JSON", "确认健康检查失败不会让前端收到不可解析响应。"),
        ("4", "ServiceClient 处理 success、patch 和失败响应", "确认内部 HTTP 客户端能正确处理成功、更新和错误路径。"),
        ("5", "knowledge-service 搜索、知识图谱、AI context 和导入校验", "确认知识库可作为 AI 上下文来源。"),
        ("6", "Gateway 转发 knowledge-service search 和 AI context", "确认前端不直接访问知识服务。"),
        ("7", "assessment-service 与 Gateway 暴露评测 CRUD 和练习历史", "确认作业、提交、题库和练习流程可通过服务和网关访问。"),
        ("8", "用户上下文读取转发请求头", "确认下游服务能识别 Gateway 注入的用户身份。"),
        ("9", "requireInternal 校验 x-edumind-internal-key", "确认内部接口有基础保护。"),
        ("10", "identity-service 登录、Token verify 和内部用户查询", "确认轻量鉴权闭环可用。"),
        ("11", "identity-service 管理用户、班级、小组和角色权限", "确认管理侧基础数据维护可用。"),
        ("12", "operations-service 导入预览、批处理任务和审计摘要", "确认运维辅助能力可用。"),
        ("13", "learning-service 课程、Dashboard、目标、任务、笔记和内部上下文", "确认学习主流程可用。"),
        ("14", "learning-service 发布协作事件失败时不回滚主业务", "确认事件失败不会影响任务/目标本地保存。"),
        ("15", "ai-service ask、plan、summarize、响应记录和 provider health", "确认基础 AI 能力和 Provider 状态可用。"),
        ("16", "StudentAiWorkflowService JSON 归一化和 fallback", "确认学生 AI 输出不稳定时仍能返回结构化结果。"),
        ("17", "TeacherAiWorkflowService JSON 归一化和 fallback", "确认教师 AI 输出不稳定时仍能返回结构化结果。"),
        ("18", "ai-service 内部 submission review 使用 mock provider", "确认评测服务可调用 AI 初评能力。"),
        ("19", "ai-service 学生 AI workflow endpoints", "确认学生 AI 六类工作流端点可用。"),
        ("20", "ai-service 教师 AI workflow endpoints", "确认教师 AI 六类工作流端点可用。"),
        ("21", "assessment-service 作业、评分、AI review、练习、错题和 Dashboard", "确认评测闭环覆盖作业到练习复盘。"),
        ("22", "collaboration-service 消息、活动、内部事件和 SSE", "确认协作与实时同步能力可用。"),
        ("23", "notification-service 与 scheduler-service 提醒、投递和偏好", "确认通知和日程提醒可用。"),
        ("24", "Gateway 转发通知和日程接口", "确认通知/日程能力通过统一入口暴露。"),
        ("25", "LM Studio endpoint normalization", "确认模型地址能稳定归一到 chat completions 路径。"),
        ("26", "analytics-service 聚合学生、课程、教师统计并检查角色", "确认统计看板和权限判断可用。"),
        ("27", "report-service 构造结构化报告和导出 payload", "确认报告服务可输出课程/学生/AI 使用等报告。"),
        ("28", "Gateway 转发 report-service API", "确认报告通过统一入口访问。"),
        ("29", "Gateway 转发 operations import、batch 和 audit API", "确认运维功能通过统一入口访问。"),
        ("30", "Gateway 转发 identity 班级管理 API", "确认班级成员管理通过统一入口访问。"),
        ("31", "Gateway 转发登录、AI、协作、SSE 和 Dashboard 聚合", "确认主要前端工作流能走完整网关路径。"),
        ("32", "Gateway 转发 analytics endpoint 并在服务可用时保持 Dashboard analytics", "确认统计服务接入 Dashboard。"),
    ], [650, 3650, 5060])

    add_para(doc, "5.3 AI 专项测试用例明细", style="Heading 2")
    add_table(doc, ["测试分组", "覆盖内容", "通过情况"], [
        ("AI schema 归一化", "每日计划、薄弱点、任务草稿、作业指导、自检、笔记整理、教师教学计划和教师草稿 action 的结构化输出。", "15/15 通过"),
        ("AI workflow", "缺少用户上下文时拒绝、解析普通 JSON 和 Markdown JSON、Provider 返回非法内容或抛错时 fallback、学生/教师各六类工作流分发。", "10/10 通过"),
        ("AI workspace", "学生结果保存、列表、详情、时间线、任务草稿确认、笔记保存；教师结果、草稿、确认动作和教师权限。", "10/10 通过"),
        ("AI 客户端适配", "学生/教师 Adapter 优先使用正式 API，接口不可用时回退到基础 AI 问答或摘要接口。", "6/6 通过"),
        ("Prompt 与仓储", "Prompt 上下文裁剪、模板渲染、endpoint 归一化、学生/教师 AI 仓储过滤排序。", "6/6 通过"),
        ("路由与鉴权", "缺少用户上下文、内部 key 错误、教师专属路由、跨用户访问、未知草稿、学生调用教师草稿等边界。", "6/6 通过"),
        ("ai-service 路由", "学生 result/draft 生命周期、note organize、教师查看学生证据、教师 draft lifecycle、provider health、submission review。", "4/4 通过"),
        ("前端选择器视图", "DOM 转义、共享视图稳定渲染、学生/教师 route-specific AI 面板和上下文。", "8/8 通过"),
    ], [1700, 5860, 1800])

    add_para(doc, "5.4 测试结论", style="Heading 2")
    add_bullets(doc, [
        "核心回归、服务集成和 AI 专项测试均通过，说明当前提交版具备可演示的基本稳定性。",
        "测试覆盖了正常流程，也覆盖了异常流程，例如 Provider 输出非法 JSON、Provider 抛错、下游协作服务不可用、内部 key 错误和跨用户访问。",
        "服务级测试输出中出现过下游不可用日志，这是测试刻意触发的容错场景；对应子测试通过，说明系统按预期返回错误或保留主业务结果。",
        "当前测试没有替代人工视觉验收。由于本机未安装 LibreOffice/soffice，Word 文档只完成结构校验，未完成渲染级页面检查。",
    ])

    output = output_path or DOCS / "项目管理文档.docx"
    doc.save(output)
    return output


def build_all():
    DOCS.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="edumind-doc-assets-") as temp_name:
        diagrams = make_diagrams(Path(temp_name))
        outputs = [
            build_requirement_doc(diagrams),
            build_design_doc(diagrams),
            build_management_doc(),
        ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    build_all()
